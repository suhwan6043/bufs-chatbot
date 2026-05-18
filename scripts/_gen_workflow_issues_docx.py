"""BUFS Chatbot 워크플로우 현 문제점 1p Word 생성.

audit (7-step) + 회귀 진단 결과 통합 — 우선순위 상위 4건만 1p에 압축.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).parent.parent / 'reports' / 'BUFS_workflow_issues_2026-05-18.docx'


def add_cell_shading(cell, color_hex):
    """Cell shading (e.g., header row)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_font(run, size=10, bold=False, color=None, name="맑은 고딕"):
    run.font.name = name
    # Set East Asian font (한글)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


doc = Document()

# Page setup: A4, 좁은 여백
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(1.8)
section.right_margin = Cm(1.8)

# Default style: 맑은 고딕 9.5pt
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(9.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
# 줄간격 1.15
pf = style.paragraph_format
pf.line_spacing = 1.15
pf.space_after = Pt(2)


# ── Title ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("BUFS 챗봇 워크플로우 — 현 문제점 보고")
set_font(run, size=15, bold=True, color="1F3864")
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
run = p.add_run("측정일 2026-05-18  ·  근거: 7-step code audit + 회귀 진단 (-14.64pp)  ·  branch: claude/nice-swartz-0c0cc4")
set_font(run, size=8.5, color="606060")
p.paragraph_format.space_after = Pt(6)


# ── 핵심 진단 요약 (1줄) ──
p = doc.add_paragraph()
run = p.add_run("핵심: ")
set_font(run, size=10, bold=True)
run = p.add_run(
    "정답률 baseline 83.54% → 68.90% (-14.64pp). 원인은 단일 코드 변경 아니라 "
    "5/6 인덱스·그래프 재인제스트 형식 변경 × KO_PROMPT v1 '원문 복사' 명령 누적. "
    "이외에 시스템 시간의 71.5%가 분류 LLM 타임아웃(rule_fallback)에 소진."
)
set_font(run, size=10)
p.paragraph_format.space_after = Pt(8)


# ── 문제 1 ──
p = doc.add_paragraph()
run = p.add_run("문제 1. 학사일정/수강신청 답변 ‑14pp 회귀 (해소 진행 중)")
set_font(run, size=11, bold=True, color="C00000")
p.paragraph_format.space_after = Pt(1)

bullets1 = [
    ("증상", "SCHEDULE 인텐트 17건 + REGISTRATION 5건 = 회귀 22건 (전체 -24건의 92%). 답변에 한글 날짜 '2026년 4월 22일' 대신 ISO '2026-04-22' 가 그대로 출력 → Contains-F1 토큰 매칭 실패."),
    ("원인", "5/6 18:31 수동 재인제스트 시 academic_graph 노드 구조 변경 (학년별 분리 + 시간 정보 + ISO 형식). _query_schedule 등 god 함수 fallback path 5곳이 _format_date 호출 없이 raw context로 LLM에 전달. KO_PROMPT v1이 '원문 그대로 복사' 명령하므로 LLM이 ISO를 그대로 답변에 채택."),
    ("권고", "audit P2-12 부분 적용 = academic_graph.py 5곳 (_schedule_to_result · 장바구니 · OCU 납부 · 학사일정 fallback · 휴/복학) context를 _format_date / _format_period로 한글 변환. 패치 적용·평가 진행 중. sanity 3건 (장바구니·수업일수 1/2선·수강신청) 한글 복원 확인."),
]
for label, body in bullets1:
    p = doc.add_paragraph()
    run = p.add_run(f"  • {label}: ")
    set_font(run, size=9.5, bold=True)
    run = p.add_run(body)
    set_font(run, size=9.5)
    p.paragraph_format.space_after = Pt(1)
doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ── 문제 2 ──
p = doc.add_paragraph()
run = p.add_run("문제 2. 분류 LLM(query_understanding) 88% TIMEOUT — 시간 71.5% 점유")
set_font(run, size=11, bold=True, color="C00000")
p.paragraph_format.space_after = Pt(1)

bullets2 = [
    ("증상", "쿼리당 평균 41.5s 중 28s가 query_understanding 단계 (1차 gemma3:4b 8s timeout → 2차 메인 LLM 20s timeout → 룰 폴백). 5/7 H100 측정 172건 중 88% TIMEOUT."),
    ("원인", "multi-task 1 (5/11 도입) 통합 LLM JSON 호출이 ollama keep_alive 만료 cold start 시 8s 안에 응답 못함. 즉 통합 모듈의 실효성 0 — 룰 폴백이 사실상 모든 분류 담당."),
    ("권고", "audit P0-3 (CONV_UNDERSTAND_TIMEOUT_SEC 8→3 default) 적용 → 추정 -6,000ms. P1-10 (FALLBACK 20→5) 적용 → 추가 -15,000ms. 단 정답률 -1~3pp 트레이드 — AB 측정 필수."),
]
for label, body in bullets2:
    p = doc.add_paragraph()
    run = p.add_run(f"  • {label}: ")
    set_font(run, size=9.5, bold=True)
    run = p.add_run(body)
    set_font(run, size=9.5)
    p.paragraph_format.space_after = Pt(1)
doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ── 문제 3 ──
p = doc.add_paragraph()
run = p.add_run("문제 3. chat.py 1,522 LOC god 4건 + stream/sync 90% 중복")
set_font(run, size=11, bold=True, color="C00000")
p.paragraph_format.space_after = Pt(1)

bullets3 = [
    ("증상", "chat_stream(CC 76, 511 LOC) + _inner_generator(CC 73, 464 LOC) + chat_sync(CC 47, 338 LOC) + _enrich_analysis(CC 49, 112 LOC). 9 stage 흐름이 stream/sync 양쪽에 약 284 LOC 중복."),
    ("원인", "초기 stream 구현 후 sync를 별도 함수로 복제. multi-task 1 통합 후에도 정리 안 됨. sync 결손 4건 누적 (① 컴포넌트 초기화 게이트 ② P4 재시도 74 LOC ③ Clarification timing 로그 ④ LLM try/except)."),
    ("권고", "audit P0-1 점진 분리 진행 중. P0-1a (Stage D 헬퍼) commit 0947bac · P0-1b (Stage F+G P4 retry) 코드 작성 완료(stash) · P0-1c (Stage I 후처리) 대기. sync 결손도 동시 해소."),
]
for label, body in bullets3:
    p = doc.add_paragraph()
    run = p.add_run(f"  • {label}: ")
    set_font(run, size=9.5, bold=True)
    run = p.add_run(body)
    set_font(run, size=9.5)
    p.paragraph_format.space_after = Pt(1)
doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ── 문제 4 (간략) ──
p = doc.add_paragraph()
run = p.add_run("문제 4. PIPELINE_TIMING 측정 결함 + 운영 OOM 위험")
set_font(run, size=11, bold=True, color="C00000")
p.paragraph_format.space_after = Pt(1)

p = doc.add_paragraph()
run = p.add_run("  • ")
set_font(run, size=9.5, bold=True)
run = p.add_run(
    "merge_ms / validate_ms 가 모든 케이스에서 0ms 일관 → 진단 데이터 신뢰성 저하. "
    "path=contact 케이스의 total=0ms 잘못 로깅. 또한 uvicorn 단일 worker 강제 명시 부재 — "
    "다중 worker 운영 시 14 싱글톤 × N = 5GB × N OOM. audit P0-2/P0-4/P2-13으로 별도 PR."
)
set_font(run, size=9.5)
p.paragraph_format.space_after = Pt(8)


# ── ROI 표 ──
p = doc.add_paragraph()
run = p.add_run("ROI 우선순위 (audit 14 PR 후보 중 핵심 4건)")
set_font(run, size=10.5, bold=True, color="1F3864")
p.paragraph_format.space_after = Pt(2)

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
headers = ["우선순위 / 작업", "latency 영향", "정답률 영향 / risk", "추정 시간"]
widths = [Cm(7.2), Cm(2.8), Cm(4.8), Cm(2.0)]
for i, (cell, text) in enumerate(zip(hdr_cells, headers)):
    cell.width = widths[i]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=9, bold=True, color="FFFFFF")
    add_cell_shading(cell, "1F3864")

rows = [
    ("P0 학사일정 ISO→한글 복원 (academic_graph 5곳)", "0", "+10~14pp 회복 기대", "0.5d"),
    ("P0 chat.py _run_pipeline 추출 (P0-1a/b/c)", "0", "0 (sync 결손 해소)", "1d"),
    ("P0 CONV_UNDERSTAND_TIMEOUT 8→3 단축", "-6,000 ms", "-1~3pp 예상 (AB 검증)", "0.5d"),
    ("P2 uvicorn workers=1 명시 + 측정 결함 fix", "0", "+ OOM 안전", "0.5d"),
]
for r_data in rows:
    row_cells = table.add_row().cells
    for i, text in enumerate(r_data):
        row_cells[i].width = widths[i]
        p = row_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_font(run, size=9)

# Footer note
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
run = p.add_run(
    "근거 보고서: reports/code_audit/AUDIT_REPORT.md · reports/regression_analysis/REPORT.md · reports/code_audit/07_checklist.md "
    "(14 PR 후보 P0/P1/P2 분류)"
)
set_font(run, size=8, color="606060")


OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(f"[saved] {OUT}")
