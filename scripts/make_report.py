# -*- coding: utf-8 -*-
"""BUFS RAG Chatbot 평가 보고서 생성 스크립트 (python-docx 사용)"""
import json, sys, io
from pathlib import Path
from collections import Counter
import datetime

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

# ─── python-docx ───────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ──────────────────────────────────────────────────────────────────────────
EVAL_PATH = Path(__file__).parent.parent / "data" / "eval" / "eval_results_20260307_182819.json"
OUT_PATH  = Path(__file__).parent.parent / "data" / "eval" / "BUFS_RAG_Evaluation_Report.docx"

# ─── 색상 정의 ──────────────────────────────────────────────────────────────
BLUE_DARK  = RGBColor(0x1F, 0x49, 0x7D)   # 헤딩 짙은 파랑
BLUE_MID   = RGBColor(0x2E, 0x74, 0xB5)   # 헤딩2 파랑
BLUE_LIGHT = RGBColor(0xBD, 0xD7, 0xEE)   # 표 헤더 배경 (hex bdd7ee)
GREEN_BG   = RGBColor(0xE2, 0xEF, 0xDA)   # 좋은 수치 배경
RED_BG     = RGBColor(0xFF, 0xE0, 0xE0)   # 나쁜 수치 배경
GRAY_BG    = RGBColor(0xF2, 0xF2, 0xF2)   # 교대색
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

# ─── 헬퍼 ──────────────────────────────────────────────────────────────────
def rgb_hex(rgb) -> str:
    """RGBColor 또는 str을 6자리 대문자 hex로 변환"""
    if isinstance(rgb, str):
        return rgb.upper().lstrip("#")
    return str(rgb).upper()  # RGBColor.__str__ returns hex like '1F497D'

def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = rgb_hex(rgb)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, sides=("top","bottom","left","right"), size="4", color="BFBFBF"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def cell_text(cell, text, bold=False, size=10, color=None, align=None):
    p = cell.paragraphs[0]
    p.clear()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "맑은 고딕"
    if color:
        run.font.color.rgb = color

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level==1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "맑은 고딕"
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = BLUE_DARK
        # 하단 경계선
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), "1F497D")
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = BLUE_MID
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = BLUE_MID
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    for run in p.runs:
        run.font.name = "맑은 고딕"
        run.font.size = Pt(10)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    run.font.size = Pt(10)
    return p

def make_header_row(table, headers, widths=None):
    row = table.rows[0]
    for i, (cell, h) in enumerate(zip(row.cells, headers)):
        set_cell_bg(cell, BLUE_DARK)
        set_cell_border(cell, color="FFFFFF")
        cell_text(cell, h, bold=True, size=10, color=WHITE,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def add_data_row(table, values, bg=WHITE, bold=False, aligns=None, size=9):
    row = table.add_row()
    for i, (cell, val) in enumerate(zip(row.cells, values)):
        set_cell_bg(cell, bg)
        set_cell_border(cell, color="BFBFBF")
        align = None
        if aligns:
            align = aligns[i]
        cell_text(cell, str(val), bold=bold, size=size, align=align)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def set_col_width(table, col_widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(col_widths_cm[i])

# ──────────────────────────────────────────────────────────────────────────
def main():
    print("JSON 로딩 중...")
    with open(EVAL_PATH, encoding="utf-8") as f:
        data = json.load(f)

    summary = data["summary"]["all"]
    results = data["results"]
    n = len(results)

    # 분류
    correct_cnt     = sum(1 for r in results if r.get("judge_correctness",0) == 1)
    wrong_items     = [r for r in results if r.get("judge_correctness",0) == 0]
    fmt_mismatch    = [r for r in results if r.get("judge_correctness",1)==1 and not r.get("contains_gt",False)]
    contains_ok     = [r for r in results if r.get("contains_gt",False)]

    # 인텐트 분포
    intent_counts = Counter(r.get("intent","?") for r in results)

    # LLM Judge 평균 (wrong items)
    wrong_rel = sum(r.get("judge_relevance",0) for r in wrong_items)
    wrong_fai = sum(r.get("judge_faithfulness",0) for r in wrong_items)

    # 응답시간 분포
    total_ms_list = [r["total_ms"] for r in results]
    gen_ms_list   = [r["generation_ms"] for r in results]
    ret_ms_list   = [r["retrieval_ms"] for r in results]

    fast = sum(1 for t in total_ms_list if t < 6000)
    mid  = sum(1 for t in total_ms_list if 6000 <= t < 10000)
    slow = sum(1 for t in total_ms_list if t >= 10000)

    # ─── Document 생성 ──────────────────────────────────────────────────────
    doc = Document()

    # 페이지 설정 (A4)
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin  = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin   = Cm(2.5)
    section.bottom_margin= Cm(2.5)

    # 기본 폰트
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)

    # ══════════════════════════════════════════════════════════════════════
    # 표지
    # ══════════════════════════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("BUFS RAG 챗봇 성능 평가 보고서")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = BLUE_DARK
    run.font.name = "맑은 고딕"

    doc.add_paragraph()
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub_p.add_run("2025학년도 2학기 학사안내 기반 RAG 시스템 LLM-as-a-Judge 평가")
    run2.font.size = Pt(14)
    run2.font.color.rgb = BLUE_MID
    run2.font.name = "맑은 고딕"

    for _ in range(2):
        doc.add_paragraph()

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = info_p.add_run(
        f"평가 일시: 2026년 3월 7일\n"
        f"평가 모델: EXAONE 3.5 7.8B (exaone3.5:7.8b)\n"
        f"임베딩 모델: BAAI/bge-m3\n"
        f"평가 데이터셋: rag_eval_dataset_100.jsonl (50개 항목)"
    )
    run3.font.size = Pt(11)
    run3.font.name = "맑은 고딕"

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 1. 평가 개요
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, "1. 평가 개요", 1)
    add_body(doc,
        "본 보고서는 부산외국어대학교(BUFS) 학사 챗봇의 RAG(Retrieval-Augmented Generation) "
        "시스템을 정량·정성적으로 평가한 결과를 담고 있습니다. "
        "평가는 2025학년도 2학기 학사안내 문서를 기반으로 구성된 50개 질의응답 데이터셋을 "
        "사용하여 수행하였으며, 자동 문자열 매칭(Contains GT)과 LLM 판단자(LLM-as-a-Judge) "
        "두 가지 방법을 결합하였습니다."
    )
    doc.add_paragraph()

    add_heading(doc, "1.1 평가 시스템 구성", 2)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(tbl, ["구성 요소", "세부 내용"])
    rows_data = [
        ("생성 모델", "EXAONE 3.5 7.8B (Ollama, Q4_K_M)"),
        ("임베딩 모델", "BAAI/bge-m3"),
        ("벡터 데이터베이스", "ChromaDB (376개 청크)"),
        ("지식 그래프", "NetworkX (48 노드, 9 엣지)"),
        ("리랭커", "비활성화 (disabled)"),
        ("컨텍스트 예산", "최대 2,048 토큰"),
        ("청크 설정", "크기 500자, 오버랩 80자"),
        ("평가 데이터셋", "50개 질의응답 (9개 인텐트 분류)"),
    ]
    for i, (k, v) in enumerate(rows_data):
        bg = GRAY_BG if i % 2 == 1 else WHITE
        add_data_row(tbl, [k, v], bg=bg, size=10)
    set_col_width(tbl, [5.5, 11.0])
    doc.add_paragraph()

    add_heading(doc, "1.2 평가 지표 정의", 2)
    metrics = [
        ("Exact Match", "생성 응답이 정답과 완전히 일치하는지 여부 (0/1)"),
        ("Contains GT", "정규화 후 정답 토큰이 생성 응답에 포함되는지 여부 (0/1)"),
        ("Hit Rate", "검색된 컨텍스트 중 정답 근거가 포함된 비율"),
        ("Judge Correctness", "LLM 판단자가 평가한 사실적 정확도 (0 또는 1)"),
        ("Judge Relevance", "LLM 판단자가 평가한 질문 관련성 (1~5점)"),
        ("Judge Faithfulness", "LLM 판단자가 평가한 컨텍스트 충실도 (1~5점)"),
        ("응답시간", "검색(Retrieval) + 생성(Generation) 시간의 합계 (ms)"),
    ]
    tbl2 = doc.add_table(rows=1, cols=2)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(tbl2, ["지표", "정의"])
    for i, (m, d) in enumerate(metrics):
        bg = GRAY_BG if i % 2 == 1 else WHITE
        add_data_row(tbl2, [m, d], bg=bg, size=10)
    set_col_width(tbl2, [4.5, 12.0])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 2. 주요 결과 요약
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, "2. 주요 결과 요약", 1)

    add_heading(doc, "2.1 핵심 성능 지표", 2)

    # KPI 표
    kpi_data = [
        ("총 평가 문항", f"{n}개", "—"),
        ("Exact Match", "0.0%", "형식 불일치로 0% (예상된 결과)"),
        ("Contains GT", f"{summary['contains_gt_%']:.1f}%  ({int(round(summary['contains_gt_%']*n/100))}/{n})", "정규화 기반 문자열 매칭"),
        ("Hit Rate", f"{summary['hit_rate_%']:.1f}%", "컨텍스트 검색 성공률"),
        ("Judge Correctness", f"{summary['judge_correctness_avg']*100:.1f}%  ({correct_cnt}/{n})", "LLM 판단자 정확도"),
        ("Judge Relevance", f"{summary['judge_relevance_avg']:.2f} / 5.00", "질문-응답 관련성"),
        ("Judge Faithfulness", f"{summary['judge_faithfulness_avg']:.2f} / 5.00", "컨텍스트 충실도"),
        ("인용 포함률", f"{summary['has_citation_%']:.1f}%", "출처 표기 비율"),
        ("불확실성 표현률", f"{summary['has_uncertainty_%']:.1f}%", "모름 표현 비율"),
        ("평균 응답 길이", f"{summary['avg_answer_length']:.1f}자", "생성 응답 평균 길이"),
    ]
    tbl3 = doc.add_table(rows=1, cols=3)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(tbl3, ["지표", "값", "비고"])
    for i, (k, v, note) in enumerate(kpi_data):
        bg = GRAY_BG if i % 2 == 1 else WHITE
        # 좋은/나쁜 수치 강조
        if "84" in v or "100" in v:
            bg = GREEN_BG
        elif "56" in v or "0.0%" in v:
            bg = RED_BG
        add_data_row(tbl3, [k, v, note], bg=bg, size=10,
                     aligns=[None, WD_ALIGN_PARAGRAPH.CENTER, None])
    set_col_width(tbl3, [4.5, 4.5, 7.5])
    doc.add_paragraph()

    add_heading(doc, "2.2 응답 시간 성능", 2)
    time_data = [
        ("평균 검색 시간 (Retrieval)", f"{summary['avg_retrieval_ms']:.1f} ms",
         "벡터 DB + 그래프 검색 합계"),
        ("평균 생성 시간 (Generation)", f"{summary['avg_generation_ms']:.1f} ms",
         "LLM 추론 시간"),
        ("평균 전체 응답 시간 (Total)", f"{summary['avg_total_ms']:.1f} ms",
         "검색 + 생성 합계"),
        ("6초 미만 응답", f"{fast}개 ({fast*100//n}%)", "빠른 응답"),
        ("6~10초 응답", f"{mid}개 ({mid*100//n}%)", "보통 응답"),
        ("10초 이상 응답", f"{slow}개 ({slow*100//n}%)", "느린 응답"),
        ("평균 검색 청크 수", f"벡터 {summary['avg_num_vector']:.1f}개 + 그래프 {summary['avg_num_graph']:.1f}개",
         "평균 9.7개 컨텍스트"),
        ("평균 상위 벡터 유사도", f"{summary['avg_top_vector_score']:.4f}", "BGE-M3 코사인 유사도"),
    ]
    tbl4 = doc.add_table(rows=1, cols=3)
    tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(tbl4, ["항목", "값", "비고"])
    for i, (k, v, note) in enumerate(time_data):
        bg = GRAY_BG if i % 2 == 1 else WHITE
        add_data_row(tbl4, [k, v, note], bg=bg, size=10,
                     aligns=[None, WD_ALIGN_PARAGRAPH.CENTER, None])
    set_col_width(tbl4, [5.0, 4.0, 7.5])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 3. 인텐트별 분석
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, "3. 인텐트별 분석", 1)
    add_body(doc,
        "데이터셋은 9개의 인텐트(의도) 유형으로 분류되어 있습니다. "
        "아래 표는 각 인텐트별 문항 수, 정확도 및 평균 점수를 보여줍니다."
    )
    doc.add_paragraph()

    # 인텐트별 통계 계산
    intent_stats = {}
    for r in results:
        intent = r.get("intent", "기타")
        if intent not in intent_stats:
            intent_stats[intent] = {"n":0, "judge_correct":0, "contains_gt":0,
                                     "relevance":[], "faithfulness":[]}
        s = intent_stats[intent]
        s["n"] += 1
        s["judge_correct"] += r.get("judge_correctness", 0)
        s["contains_gt"]   += 1 if r.get("contains_gt", False) else 0
        if r.get("judge_relevance") is not None:
            s["relevance"].append(r["judge_relevance"])
        if r.get("judge_faithfulness") is not None:
            s["faithfulness"].append(r["judge_faithfulness"])

    tbl5 = doc.add_table(rows=1, cols=6)
    tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(tbl5, ["인텐트", "문항수", "Judge\n정확도", "Contains\nGT",
                            "평균\n관련성", "평균\n충실도"])
    intent_order = sorted(intent_stats.keys(), key=lambda x: -intent_stats[x]["n"])
    for i, intent in enumerate(intent_order):
        s = intent_stats[intent]
        jc  = f"{s['judge_correct']/s['n']*100:.0f}%"
        cgt = f"{s['contains_gt']/s['n']*100:.0f}%"
        rel = f"{sum(s['relevance'])/len(s['relevance']):.2f}" if s["relevance"] else "—"
        fai = f"{sum(s['faithfulness'])/len(s['faithfulness']):.2f}" if s["faithfulness"] else "—"
        bg = GRAY_BG if i % 2 == 1 else WHITE
        add_data_row(tbl5, [intent, str(s["n"]), jc, cgt, rel, fai], bg=bg, size=10,
                     aligns=[None, WD_ALIGN_PARAGRAPH.CENTER]*3)
    set_col_width(tbl5, [4.0, 1.5, 2.0, 2.0, 2.0, 2.0])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 4. Contains GT vs Judge Correctness 비교 분석
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, "4. Contains GT vs LLM Judge 비교 분석", 1)
    add_body(doc,
        "두 평가 지표 간에 28%p의 차이(Contains GT 56% vs Judge Correctness 84%)가 관찰되었습니다. "
        "이 차이는 세 가지 유형의 케이스로 분해됩니다."
    )
    doc.add_paragraph()

    # 분류 표
    cat_data = [
        ("Judge=정답, Contains GT=일치", f"{int(summary['contains_gt_%']*n/100)}개 ({summary['contains_gt_%']:.0f}%)",
         "완전한 정답 (두 지표 모두 통과)"),
        ("Judge=정답, Contains GT=불일치", f"{len(fmt_mismatch)}개 ({len(fmt_mismatch)*100//n}%)",
         "표현 형식 차이 (실질적으로는 정답)"),
        ("Judge=오답", f"{len(wrong_items)}개 ({len(wrong_items)*100//n}%)",
         "실제 오답 케이스"),
        ("합계", f"{n}개 (100%)", "전체 평가 문항"),
    ]
    tbl6 = doc.add_table(rows=1, cols=3)
    tbl6.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(tbl6, ["분류", "문항 수", "설명"])
    bgs = [GREEN_BG, RGBColor(0xFF,0xF2,0xCC), RED_BG, GRAY_BG]
    for i, (k, v, note) in enumerate(cat_data):
        add_data_row(tbl6, [k, v, note], bg=bgs[i], size=10,
                     aligns=[None, WD_ALIGN_PARAGRAPH.CENTER, None])
    set_col_width(tbl6, [5.5, 3.0, 8.0])
    doc.add_paragraph()

    add_heading(doc, "4.1 형식 미스매치 분석", 2)
    add_body(doc,
        f"Judge=정답이지만 Contains GT=불일치인 케이스는 총 {len(fmt_mismatch)}개입니다. "
        "챗봇이 정확한 정보를 제공했지만 표현 방식이 달라 자동 매칭에서 실패한 경우입니다."
    )
    doc.add_paragraph()

    tbl7 = doc.add_table(rows=1, cols=4)
    tbl7.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(tbl7, ["질문 ID", "질문 (요약)", "정답", "챗봇 응답 (요약)"])
    for i, r in enumerate(fmt_mismatch):
        q_short = r["question"][:25] + ("…" if len(r["question"])>25 else "")
        gt_short = r["ground_truth"][:25] + ("…" if len(r["ground_truth"])>25 else "")
        ans_short = r["answer"][:35] + ("…" if len(r["answer"])>35 else "")
        bg = GRAY_BG if i % 2 == 1 else WHITE
        add_data_row(tbl7, [r["id"], q_short, gt_short, ans_short], bg=bg, size=9)
    set_col_width(tbl7, [1.5, 4.5, 4.0, 6.5])
    doc.add_paragraph()

    add_body(doc, "형식 미스매치의 주요 패턴:")
    patterns = [
        "날짜 형식 차이: '2025년 9월 3일~5일' vs '9月 3일(수) ~ 9月 5일(금)' 등",
        "URL 형식: 'http://' 포함 여부, 대소문자 차이",
        "단위/표현: '09:45' vs '9시 45분', 'A' vs 'A이다' 등",
        "정보량 차이: 정답보다 상세한 설명 포함으로 핵심 토큰 매칭 실패",
    ]
    for pt in patterns:
        add_bullet(doc, pt)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 5. 실패 케이스 상세 분석
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, "5. 실제 오답 케이스 상세 분석", 1)
    add_body(doc,
        f"Judge Correctness=0으로 평가된 실제 오답은 총 {len(wrong_items)}개입니다. "
        "아래 표에서 각 케이스의 질문, 정답, 챗봇 응답, LLM 판단 점수를 확인할 수 있습니다."
    )
    doc.add_paragraph()

    # 오답 케이스 표
    tbl8 = doc.add_table(rows=1, cols=5)
    tbl8.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(tbl8, ["ID", "질문", "정답 (GT)", "챗봇 응답", "점수\n(R/F)"])
    for i, r in enumerate(wrong_items):
        q  = r["question"][:30] + ("…" if len(r["question"])>30 else "")
        gt = r["ground_truth"][:30] + ("…" if len(r["ground_truth"])>30 else "")
        ans= r["answer"][:40] + ("…" if len(r["answer"])>40 else "")
        rel = r.get("judge_relevance", "—")
        fai = r.get("judge_faithfulness", "—")
        score = f"{rel}/{fai}"
        bg = RED_BG if i % 2 == 0 else RGBColor(0xFF, 0xE8, 0xE8)
        add_data_row(tbl8, [r["id"], q, gt, ans, score], bg=bg, size=9)
    set_col_width(tbl8, [1.2, 4.0, 3.5, 5.5, 1.5])
    doc.add_paragraph()

    add_heading(doc, "5.1 오답 원인 분류", 2)
    fail_cats = [
        ("컨텍스트 미포함 (정보 부재)", "3개",
         "q019, q020, q024: 야간수업 교시표, OCU 개강 시간 등 그래프에 저장되지 않은 정보"),
        ("잘못된 수치 생성 (환각)", "3개",
         "q015(21학점→22학점 오류), q037(30학점→학번별 세분화 혼동), 기타"),
        ("정보 검색 실패", "2개",
         "q011: 2019학번 재수강 기준 정보가 검색되지 않아 '정보 없음' 응답"),
    ]
    tbl9 = doc.add_table(rows=1, cols=3)
    tbl9.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(tbl9, ["원인 유형", "건수", "상세"])
    for i, (k, v, d) in enumerate(fail_cats):
        bg = GRAY_BG if i % 2 == 1 else WHITE
        add_data_row(tbl9, [k, v, d], bg=bg, size=10)
    set_col_width(tbl9, [4.0, 1.5, 11.0])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 6. LLM-as-a-Judge 방법론 및 신뢰도
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, "6. LLM-as-a-Judge 방법론", 1)

    add_heading(doc, "6.1 평가 설계", 2)
    add_body(doc, "LLM 판단자(Judge)는 동일한 EXAONE 3.5 7.8B 모델을 사용하였으며, "
             "다음 세 가지 차원을 독립적으로 평가합니다:")
    judge_dims = [
        ("Correctness (정확도)", "이진 점수 (0 또는 1)", "골든 컨텍스트와 생성 답변 비교"),
        ("Relevance (관련성)", "5점 척도 (1~5)", "질문과 생성 답변의 관련성 측정"),
        ("Faithfulness (충실도)", "5점 척도 (1~5)", "검색된 컨텍스트에 대한 충실성 측정"),
    ]
    tbl10 = doc.add_table(rows=1, cols=3)
    tbl10.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(tbl10, ["차원", "평가 척도", "평가 기준"])
    for i, (d, s, c) in enumerate(judge_dims):
        bg = GRAY_BG if i % 2 == 1 else WHITE
        add_data_row(tbl10, [d, s, c], bg=bg, size=10)
    set_col_width(tbl10, [5.0, 3.0, 8.5])
    doc.add_paragraph()

    add_heading(doc, "6.2 Judge 점수 해석", 2)
    add_body(doc,
        "Judge Correctness 84%는 문자열 매칭 기반의 Contains GT 56%보다 "
        "28%p 높게 나타났습니다. 이는 자동 매칭의 한계(형식 차이, 표현 다양성)를 "
        "LLM 판단자가 보완하고 있음을 보여줍니다."
    )
    doc.add_paragraph()

    judge_summary = [
        ("Judge Correctness", f"{summary['judge_correctness_avg']*100:.1f}%", "우수"),
        ("Judge Relevance", f"{summary['judge_relevance_avg']:.2f}/5.00", "매우 우수"),
        ("Judge Faithfulness", f"{summary['judge_faithfulness_avg']:.2f}/5.00", "양호"),
    ]
    tbl11 = doc.add_table(rows=1, cols=3)
    tbl11.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(tbl11, ["Judge 지표", "전체 평균", "평가"])
    bgs2 = [GREEN_BG, GREEN_BG, GRAY_BG]
    for i, (k, v, note) in enumerate(judge_summary):
        add_data_row(tbl11, [k, v, note], bg=bgs2[i], size=10,
                     aligns=[None, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER])
    set_col_width(tbl11, [5.0, 4.0, 7.5])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 7. 개선 방향
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, "7. 개선 방향 및 권고사항", 1)

    improvements = [
        ("지식 그래프 확장",
         "야간수업 교시표, OCU 관련 세부 정보 등 현재 그래프에 없는 정보를 노드/엣지로 추가하여 "
         "그래프 검색 커버리지를 높입니다. (q019, q020, q024 오답 해결)"),
        ("청킹 전략 개선",
         "표 데이터(수업 교시표 등)의 청킹 방식을 개선하여 구조화된 수치 정보가 올바르게 "
         "검색되도록 합니다. 현재 벡터 검색 평균 유사도 0.69는 개선 여지가 있습니다."),
        ("리랭커 활성화",
         "현재 disabled 상태인 리랭커를 활성화하면 관련성이 낮은 청크가 필터링되어 "
         "Judge Faithfulness(현재 4.06/5)가 개선될 것으로 예상됩니다."),
        ("학번별 조건 처리 강화",
         "q011(2019학번 재수강 기준), q037(장바구니 최대 학점) 등 학번별 조건이 복잡한 "
         "질문에서 환각이 발생합니다. 학번 분기 로직을 그래프 엣지로 명시적으로 모델링합니다."),
        ("평가 지표 개선",
         "Contains GT 정규화 로직을 개선하여 형식 차이(날짜 표현, URL 형식 등)를 "
         "허용하면 실제 정확도를 더 정확하게 측정할 수 있습니다."),
        ("응답 시간 최적화",
         f"평균 응답 시간 {summary['avg_total_ms']/1000:.1f}초 중 생성 시간이 "
         f"{summary['avg_generation_ms']/1000:.1f}초({summary['avg_generation_ms']/summary['avg_total_ms']*100:.0f}%)를 차지합니다. "
         "컨텍스트 압축 또는 모델 경량화를 통해 응답 속도를 개선할 수 있습니다."),
    ]

    tbl12 = doc.add_table(rows=1, cols=2)
    tbl12.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(tbl12, ["개선 항목", "상세 내용"])
    for i, (title, detail) in enumerate(improvements):
        bg = GRAY_BG if i % 2 == 1 else WHITE
        add_data_row(tbl12, [title, detail], bg=bg, size=10)
    set_col_width(tbl12, [4.5, 12.0])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 8. 결론
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, "8. 결론", 1)
    add_body(doc,
        "BUFS RAG 챗봇은 2025학년도 2학기 학사안내를 기반으로 한 50개 질의응답 평가에서 "
        "다음과 같은 성능을 기록하였습니다."
    )
    doc.add_paragraph()

    conclusions = [
        ("강점",
         f"• Hit Rate 100%: 모든 질문에 대해 관련 컨텍스트를 검색하는 데 성공\n"
         f"• Judge Correctness 84%: LLM 판단자 기준 높은 정확도\n"
         f"• Judge Relevance 4.66/5: 질문-응답 관련성 매우 높음\n"
         f"• 출처 표기 100%: 모든 응답에 출처 포함\n"
         f"• 인용 기반 응답으로 할루시네이션 최소화"),
        ("약점",
         f"• Contains GT 56%: 형식 불일치로 자동 매칭 낮음\n"
         f"• 실제 오답 8개(16%): 주로 그래프 미포함 정보, 복잡한 조건 처리 오류\n"
         f"• 평균 응답 시간 {summary['avg_total_ms']/1000:.1f}초: 실사용에서 개선 필요\n"
         f"• 야간수업 교시 등 일부 정보 접근 불가"),
        ("종합 평가",
         "전반적으로 학사 안내 챗봇으로서 실용적인 수준의 성능을 달성하였습니다. "
         "지식 그래프 확장과 리랭커 활성화를 통해 Judge Correctness 90% 이상 달성이 "
         "가능할 것으로 예상됩니다."),
    ]
    tbl13 = doc.add_table(rows=1, cols=2)
    tbl13.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header_row(tbl13, ["구분", "내용"])
    cats_bg = [GREEN_BG, RED_BG, GRAY_BG]
    for i, (k, v) in enumerate(conclusions):
        row = tbl13.add_row()
        set_cell_bg(row.cells[0], cats_bg[i])
        set_cell_border(row.cells[0], color="BFBFBF")
        cell_text(row.cells[0], k, bold=True, size=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        set_cell_bg(row.cells[1], cats_bg[i])
        set_cell_border(row.cells[1], color="BFBFBF")
        p = row.cells[1].paragraphs[0]
        p.clear()
        run = p.add_run(v)
        run.font.name = "맑은 고딕"
        run.font.size = Pt(10)
        row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    set_col_width(tbl13, [2.5, 14.0])
    doc.add_paragraph()

    # 푸터: 생성 날짜
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer_p.add_run(f"보고서 생성일: {datetime.date.today().strftime('%Y년 %m월 %d일')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.font.name = "맑은 고딕"

    # ─── 저장 ──────────────────────────────────────────────────────────────
    doc.save(str(OUT_PATH))
    print(f"\n✅ 보고서 저장 완료: {OUT_PATH}")
    return str(OUT_PATH)

if __name__ == "__main__":
    main()
