"""
첨부파일 인제스트 전체 파이프라인 테스트
실행: .venv/Scripts/python scripts/test_full_pipeline.py

테스트 순서:
  1. 크롤러로 공지사항 수집
  2. 첨부파일이 있는 공지 선별
  3. 첨부파일 다운로드
  4. 각 파일 형식별 추출기로 텍스트 추출
  5. ChromaDB 삽입 (임시 DB)
  6. 삽입된 청크 조회 확인
"""

import sys
import logging
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s - %(message)s",
)
logger = logging.getLogger("test_pipeline")

# ──────────────────────────────────────────────
# 1. 크롤링
# ──────────────────────────────────────────────
print("\n[1/6] 공지사항 크롤링...")
from app.crawler.notice_crawler import NoticeCrawler
crawler = NoticeCrawler()
items = crawler.crawl()
print(f"  => 수집: {len(items)}건")
for it in items:
    att_count = len(it.attachments)
    print(f"     - {it.title[:40]}  [첨부 {att_count}개]  {it.metadata.get('post_date','')}")

# 첨부파일 있는 공지만 선별
items_with_attach = [it for it in items if it.attachments]
print(f"\n  첨부파일 있는 공지: {len(items_with_attach)}건")
if not items_with_attach:
    print("  => 첨부파일 없음, 추출기 단독 테스트로 전환\n")

# ──────────────────────────────────────────────
# 2. 추출기 단독 테스트 (합성 파일)
# ──────────────────────────────────────────────
print("\n[2/6] 추출기 단독 테스트 (합성 파일 생성)...")

import io, struct, zlib

# ── DOCX 합성 ──
from docx import Document as DocxDocument
tmp_dir = Path(tempfile.mkdtemp())
docx_path = tmp_dir / "test.docx"
doc = DocxDocument()
doc.add_heading("2026학년도 1학기 수강신청 안내", 0)
doc.add_paragraph("2024학번 이후 학생은 아래 일정에 따라 수강신청을 진행하시기 바랍니다.")
doc.add_paragraph("수강신청 기간: 2026년 3월 2일(월) ~ 3월 6일(금)")
tbl = doc.add_table(rows=3, cols=2)
tbl.rows[0].cells[0].text = "구분"
tbl.rows[0].cells[1].text = "일정"
tbl.rows[1].cells[0].text = "재학생"
tbl.rows[1].cells[1].text = "3월 2일 09:00"
tbl.rows[2].cells[0].text = "신입생"
tbl.rows[2].cells[1].text = "3월 4일 09:00"
doc.add_paragraph("문의: 교학처 (051-000-0000)")
doc.save(str(docx_path))

from app.ingestion.docx_extractor import DocxExtractor
docx_pages = DocxExtractor().extract(str(docx_path))
print(f"  DOCX => {len(docx_pages)}섹션, 텍스트={sum(len(p.text) for p in docx_pages)}자, 테이블={sum(len(p.tables) for p in docx_pages)}개")
if docx_pages:
    print(f"    텍스트 미리보기: {docx_pages[0].text[:80]!r}")
    if docx_pages[0].tables:
        print(f"    테이블 미리보기:\n{docx_pages[0].tables[0][:200]}")

# ── XLSX 합성 ──
from openpyxl import Workbook
xlsx_path = tmp_dir / "test.xlsx"
wb = Workbook()
ws = wb.active
ws.title = "수강신청 현황"
ws.append(["학번", "이름", "학과", "신청 교과목", "학점"])
ws.append(["20240001", "홍길동", "영어학과", "영어회화1", 3])
ws.append(["20240002", "김철수", "일본어학과", "일본어문법", 3])
ws.append(["20240003", "이영희", "중국어학과", "중국어회화", 2])
ws2 = wb.create_sheet("졸업요건")
ws2.append(["학번", "필수학점", "선택학점", "총학점"])
ws2.append(["2024학번 이후", 30, 90, 120])
wb.save(str(xlsx_path))

from app.ingestion.xlsx_extractor import XlsxExtractor
xlsx_pages = XlsxExtractor().extract(str(xlsx_path))
print(f"  XLSX => {len(xlsx_pages)}시트")
for p in xlsx_pages:
    print(f"    시트{p.page_number}: 텍스트={p.text!r}  테이블={len(p.tables)}개")
    if p.tables:
        print(f"    테이블:\n{p.tables[0][:300]}")

# ── HWPX 합성 ──
# HWPX = ZIP(Contents/section0.xml)
import zipfile
hwpx_path = tmp_dir / "test.hwpx"
section_xml = """<?xml version="1.0" encoding="UTF-8"?>
<hh:sec xmlns:hh="http://www.hancom.co.kr/hwpml/2012/section">
  <hh:p><hh:run><hh:t>2025\ud559\ub144\ub3c4 2\ud559\uae30 \uc878\uc5c5\uc608\uc815\uc790 \uc548\ub0b4</hh:t></hh:run></hh:p>
  <hh:p><hh:run><hh:t>2022\ud559\ubc88 \uc774\uc804 \ud559\uc0dd\uc740 \uc878\uc5c5\uc2ec\uc0ac \uc2e0\uccad \uae30\uac04 \ub0b4 \uc2e0\uccad\ud558\uc2dc\uae30 \ubc14\ub78d\ub2c8\ub2e4.</hh:t></hh:run></hh:p>
  <hh:p><hh:run><hh:t>\uc2e0\uccad \uae30\uac04: 2025\ub144 10\uc6d4 1\uc77c(\uc218) ~ 10\uc6d4 10\uc77c(\uae08)</hh:t></hh:run></hh:p>
</hh:sec>""".encode("utf-8")
with zipfile.ZipFile(str(hwpx_path), "w") as zf:
    zf.writestr("Contents/section0.xml", section_xml)
    zf.writestr("mimetype", "application/x-hwp+zip")

from app.ingestion.hwp_extractor import HwpExtractor
hwpx_pages = HwpExtractor().extract(str(hwpx_path))
print(f"  HWPX => {len(hwpx_pages)}섹션, 텍스트={sum(len(p.text) for p in hwpx_pages)}자")
if hwpx_pages:
    print(f"    텍스트 미리보기: {hwpx_pages[0].text[:120]!r}")

# ──────────────────────────────────────────────
# 3. 청킹 테스트
# ──────────────────────────────────────────────
print("\n[3/6] 청킹 (pages_to_chunks)...")
from app.ingestion.chunking import pages_to_chunks

common_meta = {
    "source_notice_url": "https://www.bufs.ac.kr/test",
    "source_url": "https://www.bufs.ac.kr/test",
    "title": "테스트 공지",
    "post_date": "2026-03-10",
    "source_name": "학사공지",
    "crawled_at": "2026-03-10T10:00:00",
    "bo_table": "notice_aca",
}

all_test_pages = [
    ("docx", str(docx_path), docx_pages, "docx"),
    ("xlsx", str(xlsx_path), xlsx_pages, "xlsx"),
    ("hwpx", str(hwpx_path), hwpx_pages, "hwpx"),
]

all_chunks = []
for label, src_file, pages, file_type in all_test_pages:
    if not pages:
        print(f"  {label.upper()}: 페이지 없음 (청킹 스킵)")
        continue
    chunks = pages_to_chunks(
        pages=pages,
        source_file=src_file,
        doc_type="notice_attachment",
        semester="2026-1",
        extra_metadata={**common_meta, "filename": Path(src_file).name, "file_type": file_type},
    )
    all_chunks.extend(chunks)
    print(f"  {label.upper()} => {len(chunks)}청크")
    for ch in chunks[:2]:
        print(f"    chunk_id={ch.chunk_id[:12]}  text={ch.text[:60]!r}...")
        print(f"    meta keys: {list(ch.metadata.keys())}")

# ──────────────────────────────────────────────
# 4. ChromaDB 삽입 (임시 DB)
# ──────────────────────────────────────────────
print(f"\n[4/6] ChromaDB 삽입 ({len(all_chunks)}청크)...")
import chromadb

tmp_chroma = tmp_dir / "chroma_test"
# ChromaStore는 settings에서 경로를 읽으므로 직접 chromadb 클라이언트 사용
client = chromadb.PersistentClient(path=str(tmp_chroma))
from app.embedding import Embedder
embedder = Embedder()
collection = client.get_or_create_collection("test_attach")

# 직접 임베딩 + 삽입
for chunk in all_chunks:
    vec = embedder.embed_passage(chunk.text)
    # 메타데이터에서 None 값 제거 (ChromaDB는 None 불허)
    meta = {k: (v if v is not None else "") for k, v in {
        **chunk.metadata,
        "cohort_from": chunk.cohort_from,
        "cohort_to": chunk.cohort_to,
        "semester": chunk.semester,
        "source_file": chunk.source_file,
    }.items()}
    collection.upsert(
        ids=[chunk.chunk_id],
        embeddings=[vec.tolist()],  # type: ignore
        documents=[chunk.text],
        metadatas=[meta],
    )
print(f"  => 삽입 완료")
print(f"  => 삽입 완료")

# ──────────────────────────────────────────────
# 5. 삽입된 청크 조회 확인
# ──────────────────────────────────────────────
print("\n[5/6] ChromaDB 조회 검증...")
result = collection.get(include=["metadatas", "documents"])
ids = result.get("ids", [])
docs = result.get("documents", [])
metas = result.get("metadatas", [])

print(f"  저장된 청크 수: {len(ids)}")
for i, (doc_id, doc, meta) in enumerate(zip(ids, docs, metas)):
    print(f"\n  [{i+1}] id={doc_id[:12]}...")
    print(f"       file_type : {meta.get('file_type','?')}")
    print(f"       filename  : {meta.get('filename','?')}")
    print(f"       post_date : {meta.get('post_date','?')}")
    print(f"       source_name: {meta.get('source_name','?')}")
    print(f"       semester  : {meta.get('semester','?')}")
    print(f"       doc_type  : {meta.get('doc_type','?')}")
    print(f"       is_table  : {meta.get('is_table', False)}")
    print(f"       text      : {doc[:80]!r}...")

# ──────────────────────────────────────────────
# 6. 실제 다운로드된 첨부파일 테스트
# ──────────────────────────────────────────────
print("\n[6/6] 실제 다운로드 파일 확인...")
from app.config import DATA_DIR
crawled_pdf_dir = DATA_DIR / "pdfs" / "crawled"
hwp_dir = DATA_DIR / "attachments" / "hwp"
other_dir = DATA_DIR / "attachments" / "other"

for label, d in [("crawled PDF", crawled_pdf_dir), ("HWP", hwp_dir), ("Other(DOCX/XLSX)", other_dir)]:
    files = list(d.glob("*")) if d.exists() else []
    print(f"  {label}: {len(files)}개 파일")
    for f in files[:3]:
        print(f"    - {f.name} ({f.stat().st_size//1024}KB)")

# 실제 다운로드된 파일이 있으면 추출기 테스트
for d, extractor, label in [
    (other_dir, DocxExtractor(), "DOCX"),
    (other_dir, XlsxExtractor(), "XLSX"),
    (hwp_dir, HwpExtractor(), "HWP"),
]:
    if not d.exists():
        continue
    for f in list(d.glob("*"))[:1]:
        ext = f.suffix.lower()
        if label == "DOCX" and ext != ".docx":
            continue
        if label == "XLSX" and ext not in (".xlsx", ".xls"):
            continue
        if label == "HWP" and ext not in (".hwp", ".hwpx"):
            continue
        print(f"\n  실제 {label} 추출 테스트: {f.name}")
        pages = extractor.extract(str(f))
        total_text = sum(len(p.text) for p in pages)
        total_tables = sum(len(p.tables) for p in pages)
        print(f"    => {len(pages)}섹션, 텍스트 {total_text}자, 테이블 {total_tables}개")
        if pages and pages[0].text:
            print(f"    미리보기: {pages[0].text[:100]!r}")

print("\n========================================")
print("전체 파이프라인 테스트 완료!")
print("========================================\n")
