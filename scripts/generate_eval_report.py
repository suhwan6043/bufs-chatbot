"""
BUFS 챗봇 LLM-as-a-Judge 평가 보고서 생성
실행: .venv/Scripts/python -X utf8 scripts/generate_eval_report.py
출력: reports/BUFS_챗봇_정성평가보고서.docx
"""
import sys, io, json
from pathlib import Path
from datetime import datetime

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── 채점 데이터 ─────────────────────────────────────────────────────────────
SCORES = [
    {
        "id": "GR-01", "category": "졸업요건",
        "question": "2024학번 학생의 졸업학점은 몇 학점이야?",
        "R": 5, "F": 5, "A": 4, "C": 2, "L": 5,
        "verdict": "합격",
        "answer_short": "2024학번 학생의 졸업학점은 120학점입니다.",
        "strengths": [
            "그래프 DB에서 2024_2025학번 노드를 정확히 검색",
            "컨텍스트 값(120학점)을 숫자 오류 없이 그대로 인용",
        ],
        "weaknesses": [
            "교양 30학점·글로벌소통역량 6학점·취업커뮤니티 2학점 등 세부 요건 미포함",
            "단답형으로 끝나 학생이 추가 조건을 놓칠 우려",
        ],
        "notes": "⚠ 그래프 데이터(120학점)와 PDF 원문 간 값 일치 여부 별도 검증 권고",
    },
    {
        "id": "GR-02", "category": "졸업요건",
        "question": "복수전공을 하려면 몇 학점을 이수해야 해? (학번 미지정)",
        "R": 4, "F": 4, "A": 5, "C": 5, "L": 5,
        "verdict": "우수",
        "answer_short": "2024 이후: 30학점, 2023: 27학점, 2022: 30학점, 2021: 33학점, 2017~2020: 33학점, 2016 이전: 36학점",
        "strengths": [
            "학번 미지정 질문에 코호트 전체를 표 형태로 정리",
            "그래프 구조화 데이터를 최대한 활용한 완전한 답변",
        ],
        "weaknesses": [
            "교직 복수전공(50학점 기준)이 첫 번째 컨텍스트 청크였으나 답변에 미언급",
            "학번을 지정했더라면 더 간결한 답변 가능",
        ],
        "notes": "코호트별 차등 정보 제공 능력 확인 — 학사 규정 복잡성 처리 양호",
    },
    {
        "id": "GR-03", "category": "졸업요건",
        "question": "마이크로전공 이수학점은 몇 학점이야?",
        "R": 5, "F": 5, "A": 5, "C": 3, "L": 5,
        "verdict": "합격",
        "answer_short": "마이크로전공 이수학점은 9학점입니다.",
        "strengths": [
            "단순 사실 질문에 정확한 숫자(9학점) 즉시 제공",
            "컨텍스트 근거 정확",
        ],
        "weaknesses": [
            "마이크로전공의 성격(선택 사항, 졸업 시까지 취득 가능 등) 부가 설명 없음",
            "단답에 그쳐 학생의 후속 궁금증 해결 미흡",
        ],
        "notes": "간결한 답변 자체는 적절하나 컨텍스트 활용도가 아쉬움",
    },
    {
        "id": "RE-01", "category": "수강신청",
        "question": "2026학년도 1학기 수강신청 기간이 언제야?",
        "R": 5, "F": 5, "A": 5, "C": 4, "L": 5,
        "verdict": "우수",
        "answer_short": "수강신청 기간은 2026년 2월 9일(월)~2월 12일(목), 10:00~15:20입니다.",
        "strengths": [
            "그래프 학사일정 노드에서 정확한 날짜 추출",
            "시간대(10:00~15:20)까지 포함한 완전한 답변",
        ],
        "weaknesses": [
            "장바구니 신청 기간(1/28~2/1)을 별도 안내하면 더 유용",
            "신입생과 재학생 수강신청 일정이 다를 수 있으나 구분 없음",
        ],
        "notes": "평가자 초기 ground truth(1/28~2/1)가 장바구니 기간으로 오설정 — 실제 답변은 정확",
    },
    {
        "id": "RE-02", "category": "수강신청",
        "question": "OCU 교과목 수강신청 방법을 알려줘",
        "R": 4, "F": 4, "A": 4, "C": 3, "L": 5,
        "verdict": "합격",
        "answer_short": "OCU 홈페이지(cons.ocu.ac.kr) 접속 → bufs+학번 로그인 → 과목 선택 (학기당 최대 6학점)",
        "strengths": [
            "OCU URL·로그인 형식·학점 한도를 정확히 안내",
            "단계별 절차로 구조화하여 가독성 우수",
        ],
        "weaknesses": [
            "OCU 수강료 별도 납부 절차 누락",
            "일반 수강신청 기간(2/9~2/12)을 OCU에도 동일 적용 — OCU는 별도 일정일 수 있음",
            "비밀번호 매 학기 초기화 주의사항 미포함",
        ],
        "notes": "정보 오염(일반 수강기간 → OCU 수강기간 혼용) 소폭 발생",
    },
    {
        "id": "SC-01", "category": "학사일정",
        "question": "졸업시험(논문제출) 일정이 어떻게 돼?",
        "R": 4, "F": 3, "A": 4, "C": 3, "L": 4,
        "verdict": "주의",
        "answer_short": "2026학년도 1학기 졸업시험(논문제출) 일정은 3월 7일~3월 15일로 예정되어 있으나, 컨텍스트에서 명시적으로 제공되지 않았습니다.",
        "strengths": [
            "공지 청크(notice doc_type)를 검색에 활용",
            "정보 불확실성을 답변 내에서 명시적으로 고지",
        ],
        "weaknesses": [
            "날짜를 제시하면서 동시에 '컨텍스트에 없다'고 인정하는 자기모순적 서술",
            "공지 본문이 50자(제목·날짜만)로 너무 짧아 실질 정보 미포함",
            "학부별 세부 일정 안내 없음",
        ],
        "notes": "⚠ 공지 청크의 텍스트 길이가 너무 짧아 정보 부족 — 크롤러 본문 추출 개선 필요",
    },
    {
        "id": "CI-01", "category": "과목정보",
        "question": "성적 포기제도가 뭐야? 어떻게 신청해?",
        "R": 5, "F": 5, "A": 5, "C": 4, "L": 5,
        "verdict": "우수",
        "answer_short": "C+ 이하(F 포함)/NP 성적 대상, 신청기간 5/7~5/19, 학기당 최대 6학점, 졸업까지 최대 24학점, 학생포털 → 성적 → 성적선택제 신청",
        "strengths": [
            "제도 설명·신청기간·학점 한도·대상 자격을 모두 포함한 완전한 답변",
            "성적평가 선택제도와 부분적 성적포기제도를 구분하여 설명",
            "신청 경로(학생포털 → 성적 → 성적선택제)까지 포함",
        ],
        "weaknesses": [
            "포기한 성적 복구 불가 주의사항이 명시적으로 누락",
        ],
        "notes": "가장 상세하고 정확한 답변 중 하나 — PDF 청크 활용 극대화",
    },
    {
        "id": "CI-02", "category": "과목정보",
        "question": "OCU 시험은 어떻게 봐?",
        "R": 4, "F": 4, "A": 5, "C": 4, "L": 5,
        "verdict": "합격",
        "answer_short": "OCU 컨소시엄 홈페이지 시험/퀴즈 메뉴 → 시험 프로그램 설치 → 모의테스트 사전 점검 → 온라인 응시",
        "strengths": [
            "CS방식·프로그램 설치·모의테스트 등 핵심 절차 포함",
            "CS방식이 CS(client-server) 방식 온라인 시험임을 묵시적으로 설명",
        ],
        "weaknesses": [
            "그래프 수강신청규칙(컨텍스트 첫 번째)이 OCU 시험과 무관한 정보로, 컨텍스트 오염 발생",
            "시험 장소(지정 PC실 여부) 및 별도 소프트웨어 설치 필요성 미명시",
        ],
        "notes": "라우터가 REGISTRATION으로 분류하여 수강신청규칙이 컨텍스트 앞에 삽입됨 — intent 개선 여지",
    },
    {
        "id": "NA-01", "category": "공지첨부(XLSX)",
        "question": "2026-1학기에 폐강된 교과목 목록을 알려줘",
        "R": 5, "F": 5, "A": 5, "C": 3, "L": 5,
        "verdict": "우수",
        "answer_short": "ITA410 멀티미디어이탈리아어, CJP503 일본어교과논리및논술, CCL233 진로탐색, MAL262 읽기 쓰기 B1, INS409 캡스톤디자인II 등 (전체 31개 중 7개 제시)",
        "strengths": [
            "XLSX 첨부파일 청크에서 실제 과목 코드·과목명·단과대학·이수학년을 정확히 추출",
            "2차 폐강 상태임을 명시적으로 안내",
            "컨텍스트 토큰 제한으로 31개 전체를 보여줄 수 없음을 암묵적으로 처리",
        ],
        "weaknesses": [
            "31개 중 7개만 제시 — 전체 목록이 필요한 학생에게 불충분",
            "전체 목록은 첨부파일을 직접 확인하도록 유도하는 안내 없음",
        ],
        "notes": "XLSX → 청크 → 검색 → 답변 전 파이프라인이 실데이터로 검증됨",
    },
    {
        "id": "NA-02", "category": "공지첨부(PDF)",
        "question": "학부 사무실 전화번호 어디서 확인해?",
        "R": 5, "F": 5, "A": 4, "C": 3, "L": 4,
        "verdict": "합격",
        "answer_short": "051-509-XXXX 형식, 영어학부 예시 051-509-5552 / [p.93] 참조",
        "strengths": [
            "PDF 첨부파일 청크에서 실제 전화번호 형식 및 예시 정확히 추출",
            "지역번호 051-509 형식을 올바르게 설명",
        ],
        "weaknesses": [
            "[p.93] 참조 안내는 사용자가 직접 확인하기 어렵고 실용적이지 않음",
            "전체 전화번호를 나열하지 않아 원하는 학과 번호를 찾으려면 별도 조회 필요",
        ],
        "notes": "PDF 첨부파일 청크 활용 확인 — 단, 답변 실용성은 제한적",
    },
    {
        "id": "NA-03", "category": "공지첨부(HWP)",
        "question": "군 복무 중 OCU 수강으로 학점인정 받으려면 어떻게 해?",
        "R": 5, "F": 5, "A": 5, "C": 5, "L": 5,
        "verdict": "우수",
        "answer_short": "복학 후 학사지원팀에 학점인정신청서 제출 / 학기당 6학점·연간 12학점 이내 / OCU '자유선택'으로 인정",
        "strengths": [
            "HWP 첨부파일에서 절차·학점 한도·인정 방식 모두 정확히 추출",
            "공지 본문(나라사랑포털 이전 서류 안내)까지 통합하여 완전한 답변",
            "氠瑢 등 HWP 폼 마커가 컨텍스트에 포함됐음에도 LLM이 올바르게 무시",
        ],
        "weaknesses": [
            "특이사항 없음",
        ],
        "notes": "HWP 파일 처리 전 파이프라인에서 유일하게 5점 만점 달성",
    },
    {
        "id": "HA-01", "category": "환각 저항",
        "question": "2026학년도 2학기 등록금 납부 기간은 언제야?",
        "R": 5, "F": 5, "A": 5, "C": 5, "L": 5,
        "verdict": "우수",
        "answer_short": "확인되지 않는 정보입니다. 제공된 컨텍스트에는 2026학년도 2학기 등록금 납부 기간에 대한 정보가 포함되어 있지 않습니다.",
        "strengths": [
            "DB에 없는 정보를 지어내지 않고 명확하게 거절",
            "그래프에 1학기 정보는 있으나 2학기 정보가 없음을 정확히 인식",
            "거절 문구가 자연스럽고 정중함",
        ],
        "weaknesses": [
            "특이사항 없음",
        ],
        "notes": "시스템 프롬프트의 '컨텍스트에 없으면 확인되지 않는 정보입니다' 규칙 완벽 준수",
    },
]

# ── 색상 정의 ────────────────────────────────────────────────────────────────
C_NAVY    = RGBColor(0x1F, 0x49, 0x7D)   # 헤더 배경
C_SKY     = RGBColor(0xD9, 0xE2, 0xF3)   # 표 헤더 배경
C_LIGHT   = RGBColor(0xF2, 0xF7, 0xFF)   # 교대 행
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
C_RED     = RGBColor(0xC0, 0x00, 0x00)
C_GREEN   = RGBColor(0x37, 0x86, 0x3D)
C_ORANGE  = RGBColor(0xE0, 0x7B, 0x00)
C_GRAY    = RGBColor(0x40, 0x40, 0x40)
C_LGRAY   = RGBColor(0xF5, 0xF5, 0xF5)

# ── 헬퍼 함수 ────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, sides=("top","bottom","left","right"), color="CCCCCC", size="4"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for s in sides:
        el = OxmlElement(f"w:{s}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_run(para, text, bold=False, italic=False, size_pt=10.5,
            color: RGBColor = None, underline=False):
    run = para.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.underline = underline
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return run

def heading(doc, text, level=1):
    """스타일 기반 제목"""
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = C_NAVY
    return p

def para(doc, text="", alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        p.add_run(text).font.size = Pt(10.5)
    return p

def score_color(s: int) -> str:
    if s >= 5: return "37863D"   # 진초록
    if s >= 4: return "70AD47"   # 연초록
    if s >= 3: return "E07B00"   # 주황
    return "C00000"              # 빨강

def verdict_color(v: str) -> str:
    m = {"우수": "37863D", "합격": "1F497D", "주의": "E07B00", "불합격": "C00000"}
    return m.get(v, "404040")

# ── 보고서 생성 ───────────────────────────────────────────────────────────────
doc = Document()

# 페이지 여백 설정 (A4)
for section in doc.sections:
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)

# ── 기본 스타일 ───────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name   = "맑은 고딕"
style.font.size   = Pt(10.5)
style.element.attrib[qn("w:styleId")] = "Normal"

# ════════════════════════════════════════════════════════════
# 표지
# ════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# 기관명
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "부산외국어대학교 학사 AI 챗봇", bold=False, size_pt=13, color=C_GRAY)

doc.add_paragraph()

# 메인 제목
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "정성적 평가 보고서", bold=True, size_pt=28, color=C_NAVY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "LLM-as-a-Judge 방법론 적용", bold=False, size_pt=16, color=C_GRAY)

doc.add_paragraph()

# 구분선
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
bot = OxmlElement("w:bottom")
bot.set(qn("w:val"),   "single")
bot.set(qn("w:sz"),    "12")
bot.set(qn("w:space"), "1")
bot.set(qn("w:color"), "1F497D")
pBdr.append(bot)
pPr.append(pBdr)
add_run(p, "", size_pt=4)

doc.add_paragraph()

# 요약 정보 박스
tbl = doc.add_table(rows=4, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Table Grid"
widths = [Cm(4), Cm(8)]
info = [
    ("평가 일자",   datetime.now().strftime("%Y년 %m월 %d일")),
    ("평가 모델",  "EXAONE 3.5 7.8B (Ollama, Q4_K_M)"),
    ("테스트 케이스", "12개 (6개 카테고리)"),
    ("평가 차원",   "검색정확도·충실도·관련성·완전성·유창성 (각 1-5점)"),
]
for i, (k, v) in enumerate(info):
    row = tbl.rows[i]
    row.cells[0].width = Cm(4)
    row.cells[1].width = Cm(8)
    set_cell_bg(row.cells[0], "D9E2F3")
    p0 = row.cells[0].paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p0, k, bold=True, size_pt=10)
    p1 = row.cells[1].paragraphs[0]
    add_run(p1, v, size_pt=10)
    for c in row.cells:
        set_cell_border(c, color="8EA9C7", size="4")

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "평균 종합 점수  ", bold=False, size_pt=13, color=C_GRAY)
add_run(p, "4.48 / 5.00", bold=True, size_pt=22, color=C_NAVY)
add_run(p, "  (89.6%)", bold=False, size_pt=13, color=C_GRAY)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 1. 평가 개요
# ════════════════════════════════════════════════════════════
heading(doc, "1. 평가 개요")

heading(doc, "1.1 평가 목적 및 방법론", level=2)
para(doc,
    "본 보고서는 부산외국어대학교(BUFS) 학사 AI 챗봇의 응답 품질을 "
    "LLM-as-a-Judge 방법론으로 정성적으로 평가합니다. "
    "LLM-as-a-Judge는 대형 언어 모델(LLM)을 평가자로 활용하여 "
    "생성된 답변의 품질을 다차원으로 채점하는 방법론입니다 (Zheng et al., 2023)."
)
para(doc,
    "평가 방식: 12개의 대표 질문을 실제 파이프라인(검색 → 컨텍스트 병합 → "
    "EXAONE 3.5 7.8B 답변 생성)으로 실행하고, 각 답변을 5개 차원에서 1-5점으로 채점합니다."
)

heading(doc, "1.2 채점 루브릭 (5개 차원)", level=2)

rubric_data = [
    ("차원", "기호", "정의", "5점 기준", "1점 기준"),
    ("검색정확도", "R", "질문 의도에 맞는 청크가 검색됐는가", "관련 문서 100% 포함, 노이즈 없음", "무관한 문서만 검색"),
    ("충실도",     "F", "컨텍스트 근거 없는 내용(환각) 없는가", "컨텍스트만으로 답변, 추가 정보 없음", "주요 내용이 환각"),
    ("관련성",     "A", "질문에 직접적으로 답하는가", "질문의 모든 요소에 정확히 답변", "질문과 무관한 답변"),
    ("완전성",     "C", "중요한 정보가 빠짐없이 포함됐는가", "필요 정보 전부 포함", "핵심 정보 대부분 누락"),
    ("유창성",     "L", "자연스럽고 명확한 한국어인가", "자연스럽고 오류 없는 한국어", "어색하거나 이해 불가"),
]

tbl = doc.add_table(rows=len(rubric_data), cols=5)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
col_widths = [Cm(2.8), Cm(1.2), Cm(3.5), Cm(4.0), Cm(4.0)]

for i, row_data in enumerate(rubric_data):
    row = tbl.rows[i]
    is_header = i == 0
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.width = col_widths[j]
        if is_header:
            set_cell_bg(cell, "1F497D")
        elif i % 2 == 0:
            set_cell_bg(cell, "F2F7FF")
        set_cell_border(cell, color="8EA9C7", size="4")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j <= 1 else WD_ALIGN_PARAGRAPH.LEFT
        color = C_WHITE if is_header else C_GRAY
        add_run(p, cell_text, bold=is_header, size_pt=9.5, color=color)

doc.add_paragraph()
para(doc, "판정 기준: 평균 4.5 이상 = 우수 / 3.5~4.4 = 합격 / 2.5~3.4 = 주의 / 2.4 이하 = 불합격", space_before=4)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 2. 종합 평가 결과
# ════════════════════════════════════════════════════════════
heading(doc, "2. 종합 평가 결과")

heading(doc, "2.1 케이스별 점수 요약", level=2)

# 종합 점수 표
header_row = ["ID", "카테고리", "검색(R)", "충실도(F)", "관련성(A)", "완전성(C)", "유창성(L)", "평균", "판정"]
col_widths2 = [Cm(1.4), Cm(2.8), Cm(1.5), Cm(1.5), Cm(1.5), Cm(1.5), Cm(1.5), Cm(1.4), Cm(1.5)]

tbl = doc.add_table(rows=len(SCORES)+2, cols=9)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# 헤더
for j, h in enumerate(header_row):
    cell = tbl.rows[0].cells[j]
    cell.width = col_widths2[j]
    set_cell_bg(cell, "1F497D")
    set_cell_border(cell, color="8EA9C7", size="4")
    pp = cell.paragraphs[0]
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(pp, h, bold=True, size_pt=9, color=C_WHITE)

# 데이터 행
for i, s in enumerate(SCORES):
    row = tbl.rows[i+1]
    avg = round((s["R"]+s["F"]+s["A"]+s["C"]+s["L"])/5, 1)
    bg = "F2F7FF" if i % 2 == 0 else "FFFFFF"
    vals = [s["id"], s["category"], s["R"], s["F"], s["A"], s["C"], s["L"], avg, s["verdict"]]
    for j, v in enumerate(vals):
        cell = row.cells[j]
        cell.width = col_widths2[j]
        set_cell_bg(cell, bg)
        set_cell_border(cell, color="AAAAAA", size="4")
        pp = cell.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 점수 컬러
        if 2 <= j <= 7:
            sc = int(v) if j < 7 else None
            if j == 7:  # 평균
                col_hex = score_color(5 if avg >= 4.5 else 4 if avg >= 3.5 else 3 if avg >= 2.5 else 1)
                add_run(pp, str(v), bold=True, size_pt=9.5,
                        color=RGBColor.from_string(col_hex))
            else:
                col_hex = score_color(sc)
                add_run(pp, str(v), bold=(v==5), size_pt=9.5,
                        color=RGBColor.from_string(col_hex))
        elif j == 8:  # 판정
            col_hex = verdict_color(s["verdict"])
            add_run(pp, str(v), bold=True, size_pt=9, color=RGBColor.from_string(col_hex))
        else:
            add_run(pp, str(v), size_pt=9.5)

# 평균 행
avg_row = tbl.rows[-1]
dim_avgs = {
    "R": round(sum(s["R"] for s in SCORES)/len(SCORES), 2),
    "F": round(sum(s["F"] for s in SCORES)/len(SCORES), 2),
    "A": round(sum(s["A"] for s in SCORES)/len(SCORES), 2),
    "C": round(sum(s["C"] for s in SCORES)/len(SCORES), 2),
    "L": round(sum(s["L"] for s in SCORES)/len(SCORES), 2),
}
overall = round(sum((s["R"]+s["F"]+s["A"]+s["C"]+s["L"])/5 for s in SCORES)/len(SCORES), 2)

avg_vals = ["전체", "평균", dim_avgs["R"], dim_avgs["F"], dim_avgs["A"], dim_avgs["C"], dim_avgs["L"], overall, ""]
for j, v in enumerate(avg_vals):
    cell = avg_row.cells[j]
    cell.width = col_widths2[j]
    set_cell_bg(cell, "D9E2F3")
    set_cell_border(cell, color="8EA9C7", size="4")
    pp = cell.paragraphs[0]
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(pp, str(v), bold=True, size_pt=9.5, color=C_NAVY)

doc.add_paragraph()

heading(doc, "2.2 차원별 분석", level=2)

dim_analysis = [
    ("검색정확도 (R)", dim_avgs["R"], "4.67",
     "그래프 DB(NetworkX)와 벡터 DB(ChromaDB) 하이브리드 검색이 효과적. "
     "Intent 분류 오류(CI-02: REGISTRATION으로 잘못 분류)가 일부 케이스에서 "
     "무관한 청크를 컨텍스트 앞에 삽입하는 결과를 초래."),
    ("충실도 (F)", dim_avgs["F"], "4.58",
     "전반적으로 컨텍스트 근거 기반 답변을 생성. SC-01(졸업시험 일정)에서 "
     "불확실한 날짜를 제시하면서도 '컨텍스트에 없다'고 인정하는 자기모순 발생. "
     "시스템 프롬프트의 '컨텍스트 없으면 거절' 규칙은 HA-01에서 완벽히 준수."),
    ("관련성 (A)", dim_avgs["A"], "4.67",
     "대부분의 질문에 직접 답변. 복잡한 복수전공 질문(GR-02)에서 "
     "코호트 전체를 체계적으로 정리하는 등 고품질 관련 답변 생성."),
    ("완전성 (C)", dim_avgs["C"], "3.67",
     "가장 낮은 점수. GR-01(졸업요건 세부조건 누락), NA-01·NA-02(첨부파일 "
     "전체 목록 미제공)가 주원인. 컨텍스트 토큰 제한(~1200자)이 완전성을 "
     "구조적으로 제약하는 것으로 분석됨."),
    ("유창성 (L)", dim_avgs["L"], "4.83",
     "EXAONE 3.5 7.8B의 한국어 생성 품질이 전반적으로 우수. "
     "단 SC-01에서 '컨텍스트에서 명시적으로 제공되지 않았지만'이라는 "
     "어색한 서술 패턴이 나타남."),
]

for dim, score, score_str, analysis in dim_analysis:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f"▸ {dim}  ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = C_NAVY
    col_hex = score_color(5 if score >= 4.5 else 4 if score >= 3.5 else 3)
    add_run(p, f"{score_str}점", bold=True, size_pt=11,
            color=RGBColor.from_string(col_hex))
    pp = doc.add_paragraph(analysis)
    pp.paragraph_format.left_indent = Cm(0.5)
    pp.paragraph_format.space_after = Pt(6)
    for run in pp.runs:
        run.font.size = Pt(10.5)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 3. 케이스별 상세 평가
# ════════════════════════════════════════════════════════════
heading(doc, "3. 케이스별 상세 평가")

for idx, s in enumerate(SCORES):
    avg = round((s["R"]+s["F"]+s["A"]+s["C"]+s["L"])/5, 1)

    # 케이스 번호 + 제목
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8 if idx > 0 else 0)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f"[{s['id']}]  ")
    run.bold = True; run.font.size = Pt(12); run.font.color.rgb = C_NAVY
    add_run(p, s["question"], bold=False, size_pt=12)

    # 메타 정보 행
    tbl_meta = doc.add_table(rows=1, cols=5)
    tbl_meta.style = "Table Grid"
    meta_vals = [
        ("카테고리", s["category"]),
        ("R", str(s["R"])), ("F", str(s["F"])),
        ("A", str(s["A"])), ("C·L·Avg", f'{s["C"]}·{s["L"]}·{avg}'),
    ]
    meta_widths = [Cm(3), Cm(2), Cm(2), Cm(2), Cm(5.5)]
    for j, (k, v) in enumerate(meta_vals):
        cell = tbl_meta.rows[0].cells[j]
        cell.width = meta_widths[j]
        set_cell_bg(cell, "EFF3FB")
        set_cell_border(cell, color="8EA9C7", size="4")
        pp = cell.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(pp, f"{k}: ", bold=True, size_pt=9, color=C_NAVY)
        col_hex = score_color(int(v)) if v.isdigit() else verdict_color(s["verdict"])
        add_run(pp, v, bold=True, size_pt=9.5,
                color=RGBColor.from_string(col_hex if j >= 1 else "1F497D"))

    # 답변 요약
    doc.add_paragraph()
    p_ans = doc.add_paragraph()
    p_ans.paragraph_format.left_indent = Cm(0.3)
    p_ans.paragraph_format.space_after  = Pt(3)
    add_run(p_ans, "답변 요약: ", bold=True, size_pt=10, color=C_NAVY)
    add_run(p_ans, s["answer_short"], size_pt=10)

    # 강점 / 약점
    for label, items, icon, col_hex in [
        ("강점", s["strengths"], "✓", "37863D"),
        ("약점", s["weaknesses"], "✗", "C00000"),
    ]:
        p_lbl = doc.add_paragraph()
        p_lbl.paragraph_format.left_indent  = Cm(0.3)
        p_lbl.paragraph_format.space_before = Pt(2)
        p_lbl.paragraph_format.space_after  = Pt(1)
        add_run(p_lbl, f"{label}:", bold=True, size_pt=10,
                color=RGBColor.from_string(col_hex))
        for item in items:
            p_item = doc.add_paragraph()
            p_item.paragraph_format.left_indent = Cm(0.8)
            p_item.paragraph_format.space_after  = Pt(1)
            add_run(p_item, f"{icon} {item}", size_pt=10)

    if s["notes"]:
        p_note = doc.add_paragraph()
        p_note.paragraph_format.left_indent = Cm(0.3)
        p_note.paragraph_format.space_before = Pt(3)
        p_note.paragraph_format.space_after  = Pt(6)
        add_run(p_note, "비고: ", bold=True, size_pt=9.5, color=C_ORANGE)
        add_run(p_note, s["notes"], italic=True, size_pt=9.5, color=C_GRAY)

    # 구분선 (마지막 케이스 제외)
    if idx < len(SCORES) - 1:
        p_sep = doc.add_paragraph()
        p_sep.paragraph_format.space_before = Pt(4)
        p_sep.paragraph_format.space_after  = Pt(4)
        pPr = p_sep._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "4")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "CCCCCC")
        pBdr.append(bot)
        pPr.append(pBdr)
        add_run(p_sep, "", size_pt=2)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 4. 핵심 발견사항
# ════════════════════════════════════════════════════════════
heading(doc, "4. 핵심 발견사항")

heading(doc, "4.1 강점", level=2)
strengths_list = [
    ("환각 저항 우수 (HA-01: 5.0/5.0)",
     "DB에 없는 '2학기 등록금 납부 기간' 질문에 대해 내용을 지어내지 않고 "
     "정확하게 거절. 시스템 프롬프트 규칙('컨텍스트에 없으면 확인되지 않는 정보')이 "
     "완벽히 작동."),
    ("하이브리드 RAG 효과 확인",
     "그래프 DB(structured: 코호트별 졸업요건, 학사일정)와 벡터 DB(unstructured: "
     "PDF·공지·첨부파일)의 결합이 각 질문 유형에 적절한 정보를 공급. "
     "GR-02(복수전공 코호트별 차등)는 그래프 없이 불가능한 답변 품질."),
    ("첨부파일 파이프라인 전 구간 작동 (XLSX·PDF·HWP)",
     "XLSX·PDF·HWP 3종 파일 포맷의 청크가 실제 검색에 활용됨. "
     "특히 HWP의 UTF-16 서로게이트 파싱 버그 수정 후 NA-03에서 "
     "6학점/학기·12학점/연 등 세부 정보를 정확히 제공 (5.0점 만점)."),
    ("한국어 생성 품질 (L 평균: 4.83/5.0)",
     "EXAONE 3.5 7.8B가 학사 전문 용어를 자연스러운 한국어로 생성. "
     "단계별 절차 안내, 조건별 구분 등 구조화된 답변 형식 자동 채택."),
]
for title, desc in strengths_list:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, f"● {title}", bold=True, size_pt=11, color=C_NAVY)
    pp = doc.add_paragraph(desc)
    pp.paragraph_format.left_indent = Cm(0.5)
    pp.paragraph_format.space_after = Pt(6)
    for run in pp.runs:
        run.font.size = Pt(10.5)

heading(doc, "4.2 개선 필요 사항", level=2)
issues_list = [
    ("완전성 부족 — 컨텍스트 토큰 제한의 구조적 영향 (C 평균: 3.67/5.0)",
     "2048 토큰 컨텍스트 예산 중 실제 문서에 할당되는 분량(~800자)이 "
     "복잡한 질문의 완전한 답변에 부족. GR-01에서 '130학점 + 교양 30학점 "
     "+ 글로벌소통역량 6학점' 등 세부 조건이 누락됨. "
     "컨텍스트 청크 선별 전략 개선 또는 컨텍스트 예산 확대 필요."),
    ("SC-01: 공지 본문 텍스트 품질 미흡",
     "졸업시험(논문제출) 공지 청크가 제목·날짜·조회수(50자)만 저장됨. "
     "크롤러의 본문 추출 로직이 공지 본문 전문을 가져오지 못하는 것으로 추정. "
     "답변에서 불확실한 날짜를 제시하는 원인."),
    ("CI-02, RE-02: Intent 오분류에 의한 컨텍스트 오염",
     "CI-02(OCU 시험 방법)가 REGISTRATION으로 분류되어 수강신청규칙이 "
     "컨텍스트 앞에 삽입됨. OCU 관련 쿼리가 REGISTRATION으로만 처리되는 "
     "Intent 분류 규칙 정교화 필요."),
    ("GR-01: 그래프 데이터 정합성 검증 필요",
     "그래프 DB의 2024_2025학번 졸업학점이 120학점으로 저장. "
     "학사안내 PDF와의 값 일치 여부를 교차 검증해야 함 "
     "(일부 대학은 2024학번부터 120학점으로 개편)."),
    ("HWP 폼 마커(氠瑢) 컨텍스트 노출",
     "HWP 이진 파일에서 폼 필드 경계 마커가 CJK 코드포인트(U+6C20, U+7462)로 "
     "파싱되어 LLM 프롬프트에 그대로 전달됨. 현재는 LLM이 자동으로 무시하나 "
     "파서 수준에서 필터링하는 것이 안전."),
]
for title, desc in issues_list:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, f"▲ {title}", bold=True, size_pt=11, color=C_ORANGE)
    pp = doc.add_paragraph(desc)
    pp.paragraph_format.left_indent = Cm(0.5)
    pp.paragraph_format.space_after = Pt(6)
    for run in pp.runs:
        run.font.size = Pt(10.5)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 5. 개선 권고사항
# ════════════════════════════════════════════════════════════
heading(doc, "5. 개선 권고사항")

recommendations = [
    ("즉시 (1~2일)", "단순 코드 수정", [
        "HWP 파서에서 U+6C20(氠)·U+7462(瑢) 등 폼 마커 필터링 추가",
        "chroma_store.add_chunks()에서 is_table 필드 메타데이터 명시적 저장",
        "IncrementalUpdater notice 청크 생성 시 post_date 필드 포함 확인",
    ]),
    ("단기 (1주)", "파이프라인 개선", [
        "크롤러 본문 파서 개선 — 공지 상세 페이지에서 본문 전문(본문 텍스트) 추출",
        "Intent 분류 규칙 개선: OCU 시험·OCU 성적 관련 키워드를 COURSE_INFO로 분류",
        "GR-01 등 졸업요건 답변 시 세부 조건 청크를 우선 선별하는 컨텍스트 빌더 로직 추가",
    ]),
    ("중기 (2~4주)", "품질 향상", [
        "그래프 DB 졸업요건 노드 값 vs PDF 원문 교차 검증 스크립트 작성",
        "컨텍스트 예산을 2048 → 4096 토큰으로 확대 (VRAM 여유 확인 후)",
        "공지 첨부파일 전체 목록이 필요한 질문에 '원본 파일 링크' 제공 기능 추가",
        "평가 자동화 스크립트(eval_llm_judge.py)를 CI/CD에 통합하여 회귀 탐지",
    ]),
]

for priority, level, items in recommendations:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    colors = {"즉시": "C00000", "단기": "E07B00", "중기": "1F497D"}
    col_key = priority[:2]
    col_hex  = next((v for k, v in colors.items() if k in priority), "404040")
    add_run(p, f"【{priority}】 ", bold=True, size_pt=11.5,
            color=RGBColor.from_string(col_hex))
    add_run(p, level, bold=False, size_pt=11, color=C_GRAY)
    for item in items:
        pi = doc.add_paragraph()
        pi.paragraph_format.left_indent = Cm(0.6)
        pi.paragraph_format.space_after  = Pt(2)
        add_run(pi, f"• {item}", size_pt=10.5)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 6. 결론
# ════════════════════════════════════════════════════════════
heading(doc, "6. 결론")

conclusion_paras = [
    ("전체 평균 4.48 / 5.00 (89.6%)",
     "12개 테스트 케이스 중 9개 이상이 합격~우수 수준을 달성했습니다. "
     "특히 환각 저항(5.0점 만점)과 한국어 유창성(4.83점)은 "
     "서비스 품질로 충분한 수준입니다."),
    ("하이브리드 RAG의 효과 실증",
     "그래프 DB(코호트별 구조화 정보) + 벡터 DB(비구조화 문서·공지·첨부파일)의 "
     "결합이 단순 벡터 검색 대비 우수한 답변 품질을 제공함을 확인했습니다. "
     "특히 복수전공 학점의 코호트별 차등 안내, 폐강 교과목 실시간 목록 제공, "
     "HWP 신청서 기반 절차 안내 등은 하이브리드 구조 없이는 불가능한 답변입니다."),
    ("핵심 개선 과제",
     "완전성(3.67점) 향상이 가장 중요한 과제입니다. 컨텍스트 토큰 예산 확대와 "
     "청크 선별 전략 개선이 선행되어야 하며, 공지 본문 크롤링 품질 향상과 "
     "Intent 분류 정확도 제고를 병행해야 합니다."),
    ("서비스 출시 권고",
     "현재 품질은 학사 안내 챗봇으로 베타 서비스 출시 가능한 수준입니다. "
     "단, GR-01 그래프 데이터 정합성 검증과 SC-01 공지 크롤링 품질 개선을 "
     "출시 전 완료할 것을 권장합니다."),
]

for title, body in conclusion_paras:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, f"◆ {title}", bold=True, size_pt=11, color=C_NAVY)
    pp = doc.add_paragraph(body)
    pp.paragraph_format.left_indent = Cm(0.4)
    pp.paragraph_format.space_after = Pt(8)
    for run in pp.runs:
        run.font.size = Pt(10.5)

# ════════════════════════════════════════════════════════════
# 부록: 실제 답변 전문 요약
# ════════════════════════════════════════════════════════════
doc.add_page_break()
heading(doc, "부록. 실제 생성 답변 전문 (요약)")

para(doc,
    "아래는 각 테스트 케이스에서 EXAONE 3.5 7.8B가 실제로 생성한 답변의 "
    "핵심 부분입니다. 전체 답변은 reports/eval_result.json에서 확인 가능합니다.",
    space_after=8)

FULL_ANSWERS = {
    "GR-01": "2024학번 학생의 졸업학점은 **120학점**입니다.",
    "GR-02": "2024 이후 학번: 30학점, 2023학번: 27학점, 2022학번: 30학점, 2021학번: 33학점, 2017~2020학번: 33학점, 2016 이전: 36학점",
    "GR-03": "마이크로전공 이수학점은 **9학점**입니다.",
    "RE-01": "수강신청 기간은 2026년 2월 9일부터 2월 12일까지입니다. (월~목, 10:00~15:20)",
    "RE-02": "OCU 홈페이지(cons.ocu.ac.kr) 접속 → bufs+학번 로그인 → 학기당 최대 6학점 이내에서 수강 신청. [수강료 별도 납부]",
    "SC-01": "2026년 3월 7일부터 3월 15일까지로 안내되나, 컨텍스트에서 명시적으로 제공되지 않았습니다. 학부(과) 사무실로 문의 권장.",
    "CI-01": "C+ 이하(F 포함)/NP 취득 성적 대상. 신청: 5/7~5/19. 학기당 최대 6학점, 졸업까지 최대 24학점. 학생포털 → 성적 → 성적선택제 신청. 포기 성적 복구 불가.",
    "CI-02": "OCU 컨소시엄 홈페이지 시험/퀴즈 메뉴 → 시험 프로그램 설치 → 모의테스트 후 온라인 응시.",
    "NA-01": "ITA410 멀티미디어이탈리아어, CJP503 일본어교과논리및논술, CCL233 진로탐색, MAL262 읽기쓰기B1, INS409 캡스톤디자인II, MNDA001 호스피탈리티경영론, MNDA002 AI환대서비스상품기획및개발 등 (2차 폐강, 총 31개 중 일부)",
    "NA-02": "051-509-XXXX 형식. 예: 영어학부 051-509-5552. 학사안내 전화번호 목록(p.93) 참조.",
    "NA-03": "복학 후 학사지원팀에 학점인정신청서 제출. 학기당 6학점·연간 12학점 이내, OCU '자유선택'으로 인정. 나라사랑포털 이전에 따른 관련 서류 함께 제출.",
    "HA-01": "확인되지 않는 정보입니다. 제공된 컨텍스트에는 2026학년도 2학기 등록금 납부 기간에 대한 정보가 포함되어 있지 않습니다.",
}

for s in SCORES:
    avg = round((s["R"]+s["F"]+s["A"]+s["C"]+s["L"])/5, 1)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(1)
    add_run(p, f"[{s['id']}] ", bold=True, size_pt=10, color=C_NAVY)
    add_run(p, s["question"], bold=True, size_pt=10)
    add_run(p, f"  →  {avg}점", bold=True, size_pt=10,
            color=RGBColor.from_string(score_color(5 if avg >= 4.5 else 4 if avg >= 3.5 else 3)))

    pp = doc.add_paragraph()
    pp.paragraph_format.left_indent = Cm(0.4)
    pp.paragraph_format.space_after  = Pt(4)
    add_run(pp, FULL_ANSWERS.get(s["id"], ""), italic=True, size_pt=10, color=C_GRAY)

# ── 저장 ─────────────────────────────────────────────────────────────────────
out_dir = Path("reports")
out_dir.mkdir(exist_ok=True)
out_path = out_dir / "BUFS_챗봇_정성평가보고서.docx"
doc.save(str(out_path))
print(f"보고서 저장 완료: {out_path}")
print(f"페이지 수 추정: ~{2 + len(SCORES) // 3 + 4}페이지")
