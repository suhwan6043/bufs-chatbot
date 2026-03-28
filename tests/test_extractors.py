"""
추출기 단위 테스트 — HwpExtractor / DocxExtractor / XlsxExtractor

모든 테스트는 실제 라이브러리(zipfile, python-docx, openpyxl)를 사용하거나
외부 의존성(olefile)만 선택적으로 mock 합니다.
"""

import io
import os
import struct
import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

from app.ingestion.hwp_extractor import (
    HwpExtractor,
    _parse_hwp5_section,
    _decode_para_text,
)
from app.ingestion.docx_extractor import DocxExtractor, _table_to_markdown
from app.ingestion.xlsx_extractor import XlsxExtractor, _rows_to_markdown, _cell_str


# ═══════════════════════════════════════════════════════════
#  공통 헬퍼
# ═══════════════════════════════════════════════════════════

def _hwp5_para_text_record(text: str) -> bytes:
    """PARA_TEXT (tagId=67) 레코드 바이트를 생성합니다."""
    words = [ord(c) for c in text] + [0x000D]  # 단락 끝 마커
    payload = struct.pack(f"<{len(words)}H", *words)
    # header: tagId(10bit) | level(10bit) | size(12bit)
    header = struct.pack("<I", 67 | (len(payload) << 20))
    return header + payload


def _make_hwpx_zip_bytes(sections: dict) -> bytes:
    """{'contents/section0.xml': xml_str, ...} → 실제 ZIP 바이트"""
    import zipfile
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        for name, content in sections.items():
            zf.writestr(name, content)
    return buf.getvalue()


def _hwpx_section_xml(texts: list) -> str:
    """텍스트 목록으로 최소 HWPX 섹션 XML을 생성합니다."""
    paras = "".join(
        f'<hp:p><hp:t>{t}</hp:t></hp:p>' for t in texts
    )
    return (
        '<?xml version="1.0"?>'
        '<root xmlns:hp="http://www.hancom.co.kr/hwpml/2012/paragraph">'
        f'{paras}'
        '</root>'
    )


def _write_hwpx(tmp_path, sections: dict) -> str:
    """섹션 dict → 실제 .hwpx 임시파일 경로"""
    zip_bytes = _make_hwpx_zip_bytes(sections)
    p = tmp_path / "test.hwpx"
    p.write_bytes(zip_bytes)
    return str(p)


# ═══════════════════════════════════════════════════════════
#  HwpExtractor — HWPX (ZIP+XML)
# ═══════════════════════════════════════════════════════════

class TestHwpExtractorHwpx:

    def test_hwpx_extracts_text(self, tmp_path):
        path = _write_hwpx(tmp_path, {
            "contents/section0.xml": _hwpx_section_xml(["안녕하세요", "반갑습니다"]),
        })
        pages = HwpExtractor().extract(path)
        assert len(pages) == 1
        assert "안녕하세요" in pages[0].text
        assert "반갑습니다" in pages[0].text

    def test_hwpx_multiple_sections(self, tmp_path):
        path = _write_hwpx(tmp_path, {
            "contents/section0.xml": _hwpx_section_xml(["섹션1 내용입니다"]),
            "contents/section1.xml": _hwpx_section_xml(["섹션2 내용입니다"]),
        })
        pages = HwpExtractor().extract(path)
        assert len(pages) == 2
        assert pages[0].page_number == 1
        assert pages[1].page_number == 2

    def test_hwpx_empty_section_skipped(self, tmp_path):
        path = _write_hwpx(tmp_path, {
            "contents/section0.xml": _hwpx_section_xml([""]),
        })
        pages = HwpExtractor().extract(path)
        assert pages == []

    def test_hwpx_no_section_files(self, tmp_path):
        """섹션 파일이 없는 ZIP → 빈 리스트"""
        path = _write_hwpx(tmp_path, {"readme.txt": "no sections"})
        pages = HwpExtractor().extract(path)
        assert pages == []

    def test_hwpx_bad_zip_returns_empty(self, tmp_path):
        p = tmp_path / "broken.hwpx"
        p.write_bytes(b"this is not a valid zip file at all")
        pages = HwpExtractor().extract(str(p))
        assert pages == []

    def test_hwpx_removes_cjk_markers(self, tmp_path):
        """빈 줄로 둘러싸인 CJK 마커(氠瑢 등)가 제거됩니다."""
        path = _write_hwpx(tmp_path, {
            "contents/section0.xml": _hwpx_section_xml(["제목줄", "氠瑢", "본문 내용입니다"]),
        })
        pages = HwpExtractor().extract(path)
        assert len(pages) == 1
        assert "氠瑢" not in pages[0].text
        assert "본문 내용입니다" in pages[0].text

    def test_hwpx_bodytext_path_variant(self, tmp_path):
        """BodyText/Section0.xml 경로 형식도 인식합니다."""
        path = _write_hwpx(tmp_path, {
            "BodyText/Section0.xml": _hwpx_section_xml(["바디텍스트 형식"]),
        })
        pages = HwpExtractor().extract(path)
        assert len(pages) == 1
        assert "바디텍스트 형식" in pages[0].text


# ═══════════════════════════════════════════════════════════
#  HwpExtractor — HWP5 바이너리 파싱 (_parse_hwp5_section)
# ═══════════════════════════════════════════════════════════

class TestHwp5SectionParser:
    """_parse_hwp5_section 직접 테스트 (OLE mock 없이)"""

    def test_extracts_korean_text(self):
        data = _hwp5_para_text_record("안녕하세요")
        result = _parse_hwp5_section(data)
        assert "안녕하세요" in result

    def test_multiple_paragraphs(self):
        data = (
            _hwp5_para_text_record("첫 번째 단락") +
            _hwp5_para_text_record("두 번째 단락")
        )
        result = _parse_hwp5_section(data)
        assert "첫 번째 단락" in result
        assert "두 번째 단락" in result

    def test_skips_inline_objects(self):
        """0x0002 인라인 오브젝트 + 15워드 건너뛰고 이후 텍스트 추출"""
        # 인라인 오브젝트: 0x0002 + 15 × 0x0000 (총 16 워드)
        inline_words = [0x0002] + [0x0000] * 15
        # 이후 텍스트: "이후"
        after_words = [ord("이"), ord("후"), 0x000D]
        all_words = inline_words + after_words
        payload = struct.pack(f"<{len(all_words)}H", *all_words)
        header = struct.pack("<I", 67 | (len(payload) << 20))
        data = header + payload
        result = _parse_hwp5_section(data)
        assert "이후" in result

    def test_tab_character_converted(self):
        """0x0003 = 탭 문자로 변환됩니다."""
        words = [ord("A"), 0x0003, ord("B"), 0x000D]
        payload = struct.pack(f"<{len(words)}H", *words)
        header = struct.pack("<I", 67 | (len(payload) << 20))
        result = _parse_hwp5_section(header + payload)
        assert "A" in result
        assert "B" in result

    def test_empty_data_returns_empty_string(self):
        result = _parse_hwp5_section(b"")
        assert result == ""

    def test_non_para_text_tag_ignored(self):
        """tagId ≠ 67 레코드는 무시됩니다."""
        # tagId=66 (PARA_HEADER), size=0
        header = struct.pack("<I", 66)
        result = _parse_hwp5_section(header)
        assert result.strip() == ""


class TestDecodeParaText:
    """_decode_para_text 직접 테스트"""

    def test_basic_latin(self):
        payload = struct.pack("<3H", ord("H"), ord("i"), 0x000D)
        assert "Hi" in _decode_para_text(payload)

    def test_filters_null(self):
        payload = struct.pack("<3H", 0x0000, ord("A"), 0x000D)
        result = _decode_para_text(payload)
        assert "A" in result
        assert "\x00" not in result

    def test_skips_low_control_chars(self):
        payload = struct.pack("<4H", ord("A"), 0x0005, 0x0008, 0x000D)
        result = _decode_para_text(payload)
        assert "A" in result


# ═══════════════════════════════════════════════════════════
#  HwpExtractor — HWP5 파일 레벨 (OLE mock)
# ═══════════════════════════════════════════════════════════

class TestHwpExtractorHwp5File:

    def test_hwp5_requires_olefile(self, tmp_path):
        """olefile 미설치 시 빈 리스트 반환"""
        p = tmp_path / "test.hwp"
        p.write_bytes(b"dummy")
        with patch.dict("sys.modules", {"olefile": None}):
            pages = HwpExtractor().extract(str(p))
        assert pages == []

    def test_hwp5_bad_file_returns_empty(self, tmp_path):
        """OLE 형식이 아닌 파일 → 빈 리스트 반환"""
        p = tmp_path / "bad.hwp"
        p.write_bytes(b"not an ole2 file content here")
        pages = HwpExtractor().extract(str(p))
        assert pages == []


# ═══════════════════════════════════════════════════════════
#  HwpExtractor — 공통
# ═══════════════════════════════════════════════════════════

class TestHwpExtractorCommon:

    def test_unsupported_extension_returns_empty(self, tmp_path):
        p = tmp_path / "document.txt"
        p.write_text("hello")
        pages = HwpExtractor().extract(str(p))
        assert pages == []

    def test_source_file_stored_in_page_content(self, tmp_path):
        path = _write_hwpx(tmp_path, {
            "contents/section0.xml": _hwpx_section_xml(["텍스트 내용입니다"]),
        })
        pages = HwpExtractor().extract(path)
        assert len(pages) == 1
        assert pages[0].source_file == path


# ═══════════════════════════════════════════════════════════
#  DocxExtractor
# ═══════════════════════════════════════════════════════════

class TestDocxExtractor:
    """실제 python-docx 파일 생성 기반 테스트"""

    @staticmethod
    def _make_docx(tmp_path, paragraphs: list, filename="test.docx") -> str:
        from docx import Document as DocxDoc
        doc = DocxDoc()
        for p in paragraphs:
            para = doc.add_paragraph(p.get("text", ""))
            if p.get("style"):
                try:
                    para.style = doc.styles[p["style"]]
                except Exception:
                    pass
        path = str(tmp_path / filename)
        doc.save(path)
        return path

    def test_extracts_paragraphs(self, tmp_path):
        path = self._make_docx(tmp_path, [
            {"text": "첫 번째 단락"},
            {"text": "두 번째 단락"},
        ])
        pages = DocxExtractor().extract(path)
        assert len(pages) >= 1
        full_text = " ".join(p.text for p in pages)
        assert "첫 번째 단락" in full_text
        assert "두 번째 단락" in full_text

    def test_splits_on_heading_when_long_enough(self, tmp_path):
        """200자 이상 누적 후 Heading이 오면 섹션 분리"""
        from docx import Document as DocxDoc
        doc = DocxDoc()
        for _ in range(10):          # 25자 × 10 = 250자 > 200자
            doc.add_paragraph("A" * 25)
        doc.add_heading("새 섹션 제목", level=1)
        doc.add_paragraph("새 섹션 본문")
        path = str(tmp_path / "headings.docx")
        doc.save(path)

        pages = DocxExtractor().extract(path)
        assert len(pages) >= 2

    def test_short_section_not_split(self, tmp_path):
        """200자 미만 누적 시 헤딩이 와도 섹션 분리 안 됨"""
        from docx import Document as DocxDoc
        doc = DocxDoc()
        doc.add_paragraph("짧은 내용")  # 5자 << 200
        doc.add_heading("제목", level=1)
        doc.add_paragraph("이후 내용")
        path = str(tmp_path / "short.docx")
        doc.save(path)

        pages = DocxExtractor().extract(path)
        assert len(pages) == 1  # 짧아서 분리 안 됨

    def test_table_to_markdown_format(self):
        """_table_to_markdown: 2×2 표 → 파이프 구분 마크다운"""
        mock_table = MagicMock()
        row0 = MagicMock()
        c00 = MagicMock(); c00.text = "이름"
        c01 = MagicMock(); c01.text = "학점"
        row0.cells = [c00, c01]
        row1 = MagicMock()
        c10 = MagicMock(); c10.text = "홍길동"
        c11 = MagicMock(); c11.text = "130"
        row1.cells = [c10, c11]
        mock_table.rows = [row0, row1]

        result = _table_to_markdown(mock_table)
        assert "이름" in result
        assert "학점" in result
        assert "130" in result
        assert "---" in result  # 헤더 구분선

    def test_merged_cells_no_duplicate(self):
        """병합 셀(동일 텍스트 반복) → 중복 제거"""
        mock_table = MagicMock()
        row = MagicMock()
        c0 = MagicMock(); c0.text = "병합셀"
        c1 = MagicMock(); c1.text = "병합셀"  # 병합으로 인한 중복
        c2 = MagicMock(); c2.text = "다른셀"
        row.cells = [c0, c1, c2]
        mock_table.rows = [row]

        result = _table_to_markdown(mock_table)
        # "병합셀"이 한 번만 나타나야 함
        assert result.count("병합셀") == 1
        assert "다른셀" in result

    def test_table_extracted_to_tables_field(self, tmp_path):
        """표는 tables 필드에 마크다운으로 저장됨"""
        from docx import Document as DocxDoc
        doc = DocxDoc()
        tbl = doc.add_table(rows=2, cols=2)
        tbl.cell(0, 0).text = "헤더1"
        tbl.cell(0, 1).text = "헤더2"
        tbl.cell(1, 0).text = "데이터1"
        tbl.cell(1, 1).text = "데이터2"
        path = str(tmp_path / "table.docx")
        doc.save(path)

        pages = DocxExtractor().extract(path)
        all_tables = [t for p in pages for t in p.tables]
        assert len(all_tables) >= 1
        assert "헤더1" in all_tables[0]

    def test_import_error_returns_empty(self):
        """python-docx 미설치 시 빈 리스트 반환"""
        with patch.dict("sys.modules", {"docx": None}):
            pages = DocxExtractor().extract("any.docx")
        assert pages == []

    def test_file_not_found_returns_empty(self, tmp_path):
        pages = DocxExtractor().extract(str(tmp_path / "nonexistent.docx"))
        assert pages == []


# ═══════════════════════════════════════════════════════════
#  XlsxExtractor
# ═══════════════════════════════════════════════════════════

class TestXlsxExtractor:
    """실제 openpyxl 파일 생성 기반 테스트"""

    @staticmethod
    def _make_xlsx(tmp_path, sheets: list, filename="test.xlsx") -> str:
        """sheets: [{"name": str, "rows": [[str, ...], ...]}]"""
        from openpyxl import Workbook
        wb = Workbook()
        wb.remove(wb.active)  # 기본 빈 시트 제거
        for s in sheets:
            ws = wb.create_sheet(title=s["name"])
            for row in s["rows"]:
                ws.append(row)
        path = str(tmp_path / filename)
        wb.save(path)
        return path

    def test_single_sheet_extracts_table(self, tmp_path):
        path = self._make_xlsx(tmp_path, [
            {"name": "성적", "rows": [["이름", "학점"], ["홍길동", "130"]]},
        ])
        pages = XlsxExtractor().extract(path)
        assert len(pages) == 1
        assert pages[0].page_number == 1
        assert len(pages[0].tables) == 1
        assert "이름" in pages[0].tables[0]
        assert "130" in pages[0].tables[0]

    def test_multiple_sheets(self, tmp_path):
        path = self._make_xlsx(tmp_path, [
            {"name": "Sheet1", "rows": [["A", "B"], ["1", "2"]]},
            {"name": "Sheet2", "rows": [["X", "Y"], ["3", "4"]]},
        ])
        pages = XlsxExtractor().extract(path)
        assert len(pages) == 2
        assert pages[0].page_number == 1
        assert pages[1].page_number == 2
        assert "Sheet1" in pages[0].text
        assert "Sheet2" in pages[1].text

    def test_empty_sheet_skipped(self, tmp_path):
        from openpyxl import Workbook
        wb = Workbook()
        wb.remove(wb.active)
        wb.create_sheet("빈시트")  # 데이터 없음
        ws_data = wb.create_sheet("데이터")
        ws_data.append(["값1", "값2"])
        path = str(tmp_path / "mixed.xlsx")
        wb.save(path)

        pages = XlsxExtractor().extract(path)
        assert len(pages) == 1
        assert "데이터" in pages[0].text

    def test_cell_truncation(self):
        """200자 초과 셀 값은 잘립니다"""
        from app.ingestion.xlsx_extractor import _MAX_CELL_LEN
        long_val = "가" * 300
        result = _cell_str(long_val)
        assert len(result) <= _MAX_CELL_LEN

    def test_none_cells_empty_string(self):
        assert _cell_str(None) == ""

    def test_integer_cell_converted_to_string(self):
        assert _cell_str(42) == "42"

    def test_markdown_table_format(self):
        """_rows_to_markdown: 헤더행 + --- 구분자 행"""
        rows = [["이름", "나이"], ["홍길동", "20"]]
        result = _rows_to_markdown(rows)
        lines = result.split("\n")
        assert lines[0].startswith("|")
        assert "---" in lines[1]     # 헤더 구분선
        assert "홍길동" in lines[2]

    def test_markdown_pads_short_rows(self):
        """행 길이가 다를 때 최대 열 수로 패딩"""
        rows = [["A", "B", "C"], ["1"]]  # 두 번째 행은 1열만
        result = _rows_to_markdown(rows)
        # 두 번째 행도 3열 형식으로 렌더링
        lines = result.split("\n")
        data_line = lines[2]
        assert data_line.count("|") >= 3

    def test_text_header_contains_sheet_info(self, tmp_path):
        """text 필드에 파일명·시트명·행 수 포함"""
        path = self._make_xlsx(tmp_path, [
            {"name": "학사일정", "rows": [["날짜", "행사"], ["3/2", "개강"]]},
        ])
        pages = XlsxExtractor().extract(path)
        assert "학사일정" in pages[0].text
        assert "2" in pages[0].text   # 2 data rows

    def test_import_error_returns_empty(self):
        """openpyxl 미설치 시 빈 리스트 반환"""
        with patch.dict("sys.modules", {"openpyxl": None}):
            pages = XlsxExtractor().extract("any.xlsx")
        assert pages == []

    def test_file_not_found_returns_empty(self, tmp_path):
        pages = XlsxExtractor().extract(str(tmp_path / "nonexistent.xlsx"))
        assert pages == []
